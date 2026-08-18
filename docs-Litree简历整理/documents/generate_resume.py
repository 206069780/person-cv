from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "web" / "src" / "data" / "resume-data.json"
DOCX_PATH = Path(__file__).with_name("付道品-高级Java开发工程师.docx")
PDF_PATH = Path(__file__).with_name("付道品-高级Java开发工程师.pdf")

FONT_REGULAR_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_REGULAR = "ResumeYaHei"
FONT_BOLD = "ResumeYaHeiBold"

NAVY = "16324A"
TEAL = "167D7A"
ORANGE = "D9772D"
INK = "17232D"
MUTED = "5D6A73"
PALE = "EAF4F3"
RULE = "CBD7DB"


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def set_run_font(run, size: float, bold: bool = False, color: str = INK) -> None:
    name = "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeatable_cell_margins(cell, top=70, start=100, bottom=70, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    end = paragraph.add_run(" / 10 页")
    set_run_font(end, 8, color=MUTED)


def configure_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(7)

    for style_name, size, before, after, color in (
        ("Heading 1", 17, 0, 8, NAVY),
        ("Heading 2", 11.5, 7, 4, TEAL),
        ("Heading 3", 9.5, 5, 2, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Cm(0.56)
        style.paragraph_format.first_line_indent = Cm(-0.28)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.1

    if "Resume Lead" not in [s.name for s in styles]:
        lead = styles.add_style("Resume Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Resume Lead"]
    lead.font.name = "Microsoft YaHei"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lead.font.size = Pt(10.5)
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(6)
    lead.paragraph_format.line_spacing = 1.2

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("付道品  |  高级 Java 开发工程师")
    set_run_font(r, 8, bold=True, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_docx_title(doc: Document, kicker: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(kicker.upper())
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 17, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(subtitle)
        set_run_font(r, 9, color=MUTED)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{label}  ")
    set_run_font(r, 8.8, bold=True, color=TEAL)
    r = p.add_run(text)
    set_run_font(r, 8.7, color=INK)


def add_docx_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run_font(r, 8.5, color=INK)


def add_stack(doc: Document, stack: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("技术栈  ")
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    r = p.add_run(" · ".join(stack))
    set_run_font(r, 8.2, color=MUTED)


def add_docx_cover(doc: Document, data: dict) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("JAVA BACKEND · AIOT · AGENT ENGINEERING")
    set_run_font(r, 9.5, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(data["profile"]["name"])
    set_run_font(r, 31, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(data["profile"]["title"])
    set_run_font(r, 17, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(data["profile"]["experience"])
    set_run_font(r, 10.5, color=MUTED)

    table = doc.add_table(rows=1, cols=len(data["highlights"]))
    table.autofit = False
    width = Cm(17.9 / len(data["highlights"]))
    for cell, value in zip(table.rows[0].cells, data["highlights"]):
        cell.width = width
        set_repeatable_cell_margins(cell, 100, 80, 100, 80)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), PALE)
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 8.2, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data["profile"]["summary"])
    set_run_font(r, 10.2, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{data['profile']['phone']}  |  {data['profile']['email']}")
    set_run_font(r, 9.5, bold=True, color=TEAL)


def add_docx_overview(doc: Document, data: dict) -> None:
    add_docx_title(doc, "02 / 10", "核心简历", "面向普通投递的能力概览与工作经历")
    p = doc.add_paragraph(style="Resume Lead")
    p.add_run(data["profile"]["summary"])

    doc.add_paragraph("核心能力", style="Heading 2")
    for strength in data["strengths"]:
        add_labeled_paragraph(doc, strength["title"], strength["evidence"])

    doc.add_paragraph("工作经历", style="Heading 2")
    for exp in data["experiences"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{exp['period']}  |  {exp['company']}  |  {exp['title']}")
        set_run_font(r, 9.1, bold=True, color=NAVY)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(exp["summary"])
        set_run_font(r, 8.4)
        add_docx_bullets(doc, exp["achievements"])

    stack = []
    for project in data["projects"]:
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)
    add_stack(doc, stack[:24])


def add_docx_topic(doc: Document, page_no: int, project: dict, topic: dict) -> None:
    add_docx_title(
        doc,
        f"{page_no:02d} / 10 · {project['company']}",
        topic["title"],
        f"{project['name']}  |  {project['period']}",
    )
    if project["id"] == "litree":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("项目事实  国内外 10,000+ 水站 · Litree 六个专题均属于同一项目")
        set_run_font(r, 8.6, bold=True, color=ORANGE)
    add_labeled_paragraph(doc, "项目定位", project["summary"])
    add_labeled_paragraph(doc, "项目背景", topic["background"])
    add_labeled_paragraph(doc, "本人角色", topic["role"])
    add_labeled_paragraph(doc, "业务链路", topic["flow"])
    add_labeled_paragraph(doc, "工程边界", topic["engineeringBoundary"])
    doc.add_paragraph("核心实现", style="Heading 2")
    add_docx_bullets(doc, topic["implementation"])
    doc.add_paragraph("技术难点", style="Heading 2")
    add_docx_bullets(doc, topic["challenges"])
    add_labeled_paragraph(doc, "落地结果", topic["outcome"])
    add_stack(doc, topic["stack"])


def build_docx(data: dict) -> None:
    doc = Document()
    configure_docx(doc)
    add_docx_cover(doc, data)
    doc.add_page_break()
    add_docx_overview(doc, data)

    pages: list[tuple[dict, dict]] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            pages.append((project, topic))
    if len(pages) != 8:
        raise ValueError(f"Expected 8 project topic pages, got {len(pages)}")
    for page_no, (project, topic) in enumerate(pages, start=3):
        doc.add_page_break()
        add_docx_topic(doc, page_no, project, topic)
    doc.save(DOCX_PATH)


def register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont(FONT_REGULAR, str(FONT_REGULAR_PATH)))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, str(FONT_BOLD_PATH)))


def pdf_text_width(text: str, font: str, size: float) -> float:
    return pdfmetrics.stringWidth(text, font, size)


def wrap_pdf_text(text: str, font: str, size: float, max_width: float) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if current and pdf_text_width(candidate, font, size) > max_width:
            lines.append(current.rstrip())
            current = char.lstrip()
        else:
            current = candidate
    if current:
        lines.append(current.rstrip())
    return lines or [""]


class PdfPage:
    def __init__(self, c: canvas.Canvas, page_no: int):
        self.c = c
        self.page_no = page_no
        self.width, self.height = A4
        self.left = 46
        self.right = self.width - 46
        self.y = self.height - 52

    def header(self, section: str) -> None:
        c = self.c
        c.setFont(FONT_BOLD, 7.8)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawString(self.left, self.height - 25, "付道品  |  高级 Java 开发工程师")
        c.drawRightString(self.right, self.height - 25, section)
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.5)
        c.line(self.left, self.height - 32, self.right, self.height - 32)

    def footer(self) -> None:
        c = self.c
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.45)
        c.line(self.left, 31, self.right, 31)
        c.setFont(FONT_REGULAR, 7.5)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawRightString(self.right, 18, f"第 {self.page_no} / 10 页")

    def title(self, kicker: str, title: str, subtitle: str | None = None) -> None:
        c = self.c
        c.setFont(FONT_BOLD, 8.5)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.drawString(self.left, self.y, kicker)
        self.y -= 22
        c.setFont(FONT_BOLD, 18)
        c.setFillColor(HexColor(f"#{NAVY}"))
        c.drawString(self.left, self.y, title)
        self.y -= 18
        if subtitle:
            c.setFont(FONT_REGULAR, 8.8)
            c.setFillColor(HexColor(f"#{MUTED}"))
            c.drawString(self.left, self.y, subtitle)
            self.y -= 14
        self.y -= 5

    def paragraph(self, text: str, size=9.2, leading=13.5, color=INK, indent=0, after=6, bold=False) -> None:
        c = self.c
        font = FONT_BOLD if bold else FONT_REGULAR
        c.setFont(font, size)
        c.setFillColor(HexColor(f"#{color}"))
        lines = wrap_pdf_text(text, font, size, self.right - self.left - indent)
        for line in lines:
            c.drawString(self.left + indent, self.y, line)
            self.y -= leading
        self.y -= after

    def labeled(self, label: str, text: str) -> None:
        c = self.c
        c.setFont(FONT_BOLD, 9.3)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.drawString(self.left, self.y, label)
        label_width = pdf_text_width(label + "  ", FONT_BOLD, 9.3)
        lines = wrap_pdf_text(text, FONT_REGULAR, 9.05, self.right - self.left - label_width)
        c.setFont(FONT_REGULAR, 9.05)
        c.setFillColor(HexColor(f"#{INK}"))
        c.drawString(self.left + label_width, self.y, lines[0])
        self.y -= 13.5
        for line in lines[1:]:
            c.drawString(self.left, self.y, line)
            self.y -= 13.5
        self.y -= 5

    def section(self, title: str) -> None:
        self.y -= 2
        self.c.setFont(FONT_BOLD, 11.3)
        self.c.setFillColor(HexColor(f"#{TEAL}"))
        self.c.drawString(self.left, self.y, title)
        self.y -= 17

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            lines = wrap_pdf_text(item, FONT_REGULAR, 9.0, self.right - self.left - 14)
            self.c.setFillColor(HexColor(f"#{ORANGE}"))
            self.c.circle(self.left + 3, self.y + 2.5, 1.7, fill=1, stroke=0)
            self.c.setFont(FONT_REGULAR, 9.0)
            self.c.setFillColor(HexColor(f"#{INK}"))
            for line in lines:
                self.c.drawString(self.left + 13, self.y, line)
                self.y -= 13.2
            self.y -= 3

    def stack(self, items: list[str]) -> None:
        self.y -= 2
        self.paragraph("技术栈  " + " · ".join(items), size=8.5, leading=12.4, color=MUTED, after=0, bold=False)


def draw_pdf_cover(c: canvas.Canvas, data: dict) -> None:
    w, h = A4
    c.setFillColor(HexColor("#F7FAFA"))
    c.rect(0, 0, w, h, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.rect(0, h - 14, w, 14, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(0, 0, w, 6, fill=1, stroke=0)

    c.setFont(FONT_BOLD, 8.8)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawCentredString(w / 2, h - 180, "JAVA BACKEND · AIOT · AGENT ENGINEERING")
    c.setFont(FONT_BOLD, 32)
    c.setFillColor(HexColor(f"#{NAVY}"))
    c.drawCentredString(w / 2, h - 235, data["profile"]["name"])
    c.setFont(FONT_BOLD, 18)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.drawCentredString(w / 2, h - 276, data["profile"]["title"])
    c.setFont(FONT_REGULAR, 10.5)
    c.setFillColor(HexColor(f"#{MUTED}"))
    c.drawCentredString(w / 2, h - 306, data["profile"]["experience"])

    x0 = 46
    total_w = w - 92
    cell_w = total_w / 4
    y = h - 386
    for i, item in enumerate(data["highlights"]):
        x = x0 + i * cell_w
        c.setFillColor(HexColor(f"#{PALE}"))
        c.roundRect(x + 3, y, cell_w - 6, 42, 3, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 8.1)
        c.setFillColor(HexColor(f"#{NAVY}"))
        lines = wrap_pdf_text(item, FONT_BOLD, 8.1, cell_w - 18)
        line_y = y + 25 + (len(lines) == 1) * -4
        for line in lines:
            c.drawCentredString(x + cell_w / 2, line_y, line)
            line_y -= 11

    c.setFont(FONT_REGULAR, 10)
    c.setFillColor(HexColor(f"#{INK}"))
    summary_lines = wrap_pdf_text(data["profile"]["summary"], FONT_REGULAR, 10, total_w - 30)
    sy = h - 488
    for line in summary_lines:
        c.drawCentredString(w / 2, sy, line)
        sy -= 15
    c.setFont(FONT_BOLD, 9.5)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.drawCentredString(w / 2, 145, f"{data['profile']['phone']}  |  {data['profile']['email']}")
    c.setFont(FONT_REGULAR, 7.5)
    c.setFillColor(HexColor(f"#{MUTED}"))
    c.drawRightString(w - 46, 18, "第 1 / 10 页")


def draw_pdf_overview(c: canvas.Canvas, data: dict) -> None:
    page = PdfPage(c, 2)
    page.header("核心简历")
    page.title("02 / 10", "核心简历", "面向普通投递的能力概览与工作经历")
    page.paragraph(data["profile"]["summary"], size=9.4, leading=13.5, color=NAVY, after=7)
    page.section("核心能力")
    for strength in data["strengths"]:
        page.labeled(strength["title"], strength["evidence"])
    page.section("工作经历")
    for exp in data["experiences"]:
        page.paragraph(f"{exp['period']}  |  {exp['company']}  |  {exp['title']}", size=8.8, leading=12, color=NAVY, after=2, bold=True)
        page.paragraph(exp["summary"], size=8.1, leading=11.2, after=1)
        page.bullets(exp["achievements"])
    stack: list[str] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)
    page.stack(stack[:24])
    page.footer()


def draw_pdf_topic(c: canvas.Canvas, page_no: int, project: dict, topic: dict) -> None:
    page = PdfPage(c, page_no)
    page.header(project["name"])
    page.title(f"{page_no:02d} / 10 · {project['company']}", topic["title"], f"{project['name']}  |  {project['period']}")
    if project["id"] == "litree":
        page.paragraph("项目事实  国内外 10,000+ 水站 · 本页属于 Litree 同一项目的专题展开", size=8.4, leading=11.5, color=ORANGE, after=6, bold=True)
    page.labeled("项目定位", project["summary"])
    page.labeled("项目背景", topic["background"])
    page.labeled("本人角色", topic["role"])
    page.labeled("业务链路", topic["flow"])
    page.labeled("工程边界", topic["engineeringBoundary"])
    page.section("核心实现")
    page.bullets(topic["implementation"])
    page.section("技术难点")
    page.bullets(topic["challenges"])
    page.labeled("落地结果", topic["outcome"])
    page.stack(topic["stack"])
    if page.y < 48:
        raise ValueError(f"PDF page {page_no} overflowed: y={page.y:.1f}")
    page.footer()


def build_pdf(data: dict) -> None:
    register_pdf_fonts()
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4, pageCompression=1)
    c.setTitle("付道品 - 高级 Java 开发工程师")
    c.setAuthor("付道品")
    c.setSubject("高级 Java 开发工程师简历")
    draw_pdf_cover(c, data)
    c.showPage()
    draw_pdf_overview(c, data)

    pages: list[tuple[dict, dict]] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            pages.append((project, topic))
    for page_no, (project, topic) in enumerate(pages, start=3):
        c.showPage()
        draw_pdf_topic(c, page_no, project, topic)
    c.save()


def main() -> None:
    data = load_data()
    build_docx(data)
    build_pdf(data)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
