from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "web" / "src" / "data" / "resume-data.json"
DOCX_PATH = ROOT / "付道品-高级Java开发工程师.docx"
PDF_PATH = ROOT / "付道品-高级Java开发工程师.pdf"

FONT_REGULAR_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_REGULAR = "ResumeYaHei"
FONT_BOLD = "ResumeYaHeiBold"

GRAPHITE = "081117"
NAVY = "12323C"
TEAL = "00A89D"
CYAN = "28D7E5"
ORANGE = "FF6B3D"
INK = "13242A"
MUTED = "62727D"
PALE = "E9F3F3"
RULE = "CBD8DC"
COLD_WHITE = "F5F8F9"


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def set_run_font(run, size: float, bold: bool = False, color: str = INK) -> None:
    name = "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeatable_cell_margins(cell, top=70, start=100, bottom=70, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    end = paragraph.add_run(" / 10 页")
    set_run_font(end, 8, color=MUTED)


def configure_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(7)

    for style_name, size, before, after, color in (
        ("Heading 1", 17, 0, 8, NAVY),
        ("Heading 2", 11.5, 7, 4, TEAL),
        ("Heading 3", 9.5, 5, 2, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Cm(0.56)
        style.paragraph_format.first_line_indent = Cm(-0.28)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.1

    if "Resume Lead" not in [s.name for s in styles]:
        lead = styles.add_style("Resume Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Resume Lead"]
    lead.font.name = "Microsoft YaHei"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lead.font.size = Pt(10.5)
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(6)
    lead.paragraph_format.line_spacing = 1.2

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("付道品  |  高级 Java 开发工程师")
    set_run_font(r, 8, bold=True, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_docx_title(doc: Document, kicker: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(kicker.upper())
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 17, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(subtitle)
        set_run_font(r, 9, color=MUTED)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{label}  ")
    set_run_font(r, 8.8, bold=True, color=TEAL)
    r = p.add_run(text)
    set_run_font(r, 8.7, color=INK)


def add_docx_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run_font(r, 8.5, color=INK)


def add_stack(doc: Document, stack: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("技术栈  ")
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    r = p.add_run(" · ".join(stack))
    set_run_font(r, 8.2, color=MUTED)


def add_docx_cover(doc: Document, data: dict) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("JAVA BACKEND · AIOT · AGENT ENGINEERING")
    set_run_font(r, 9.5, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(data["profile"]["name"])
    set_run_font(r, 31, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(data["profile"]["title"])
    set_run_font(r, 17, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(data["profile"]["experience"])
    set_run_font(r, 10.5, color=MUTED)

    table = doc.add_table(rows=1, cols=len(data["highlights"]))
    table.autofit = False
    width = Cm(17.9 / len(data["highlights"]))
    for cell, value in zip(table.rows[0].cells, data["highlights"]):
        cell.width = width
        set_repeatable_cell_margins(cell, 100, 80, 100, 80)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), PALE)
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 8.2, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data["profile"]["summary"])
    set_run_font(r, 10.2, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{data['profile']['phone']}  |  {data['profile']['email']}")
    set_run_font(r, 9.5, bold=True, color=TEAL)


def add_docx_overview(doc: Document, data: dict) -> None:
    add_docx_title(doc, "02 / 10", "核心简历", "面向普通投递的能力概览与工作经历")
    p = doc.add_paragraph(style="Resume Lead")
    p.add_run(data["profile"]["summary"])

    doc.add_paragraph("核心能力", style="Heading 2")
    for strength in data["strengths"]:
        add_labeled_paragraph(doc, strength["title"], strength["evidence"])

    doc.add_paragraph("工作经历", style="Heading 2")
    for exp in data["experiences"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{exp['period']}  |  {exp['company']}  |  {exp['title']}")
        set_run_font(r, 9.1, bold=True, color=NAVY)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(exp["summary"])
        set_run_font(r, 8.4)
        add_docx_bullets(doc, exp["achievements"])

    stack = []
    for project in data["projects"]:
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)
    add_stack(doc, stack[:24])


def add_docx_topic(doc: Document, page_no: int, project: dict, topic: dict) -> None:
    add_docx_title(
        doc,
        f"{page_no:02d} / 10 · {project['company']}",
        topic["title"],
        f"{project['name']}  |  {project['period']}",
    )
    if project["id"] == "litree":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("项目事实  国内外 10w+ 水站 · Litree 六个专题均属于同一项目")
        set_run_font(r, 8.6, bold=True, color=ORANGE)
    add_labeled_paragraph(doc, "项目定位", project["summary"])
    add_labeled_paragraph(doc, "项目背景", topic["background"])
    add_labeled_paragraph(doc, "本人角色", topic["role"])
    add_labeled_paragraph(doc, "业务链路", topic["flow"])
    add_labeled_paragraph(doc, "工程边界", topic["engineeringBoundary"])
    doc.add_paragraph("核心实现", style="Heading 2")
    add_docx_bullets(doc, topic["implementation"])
    doc.add_paragraph("技术难点", style="Heading 2")
    add_docx_bullets(doc, topic["challenges"])
    add_labeled_paragraph(doc, "落地结果", topic["outcome"])
    add_stack(doc, topic["stack"])


def build_docx(data: dict) -> None:
    doc = Document()
    configure_docx(doc)
    add_docx_cover(doc, data)
    doc.add_page_break()
    add_docx_overview(doc, data)

    pages: list[tuple[dict, dict]] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            pages.append((project, topic))
    if len(pages) != 8:
        raise ValueError(f"Expected 8 project topic pages, got {len(pages)}")
    for page_no, (project, topic) in enumerate(pages, start=3):
        doc.add_page_break()
        add_docx_topic(doc, page_no, project, topic)
    doc.save(DOCX_PATH)


def register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont(FONT_REGULAR, str(FONT_REGULAR_PATH)))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, str(FONT_BOLD_PATH)))


def pdf_text_width(text: str, font: str, size: float) -> float:
    return pdfmetrics.stringWidth(text, font, size)


def wrap_pdf_text(text: str, font: str, size: float, max_width: float) -> list[str]:
    no_line_start = set("，。！？；：、）》】〕」』”’％%)]},.;:!?")
    no_line_end = set("（《【〔「『“‘([{")
    tokens = re.findall(
        r"[A-Za-z0-9][A-Za-z0-9_./:+#-]*(?:[ \t]+[A-Za-z0-9][A-Za-z0-9_./:+#-]*)*|[ \t]+|\n|.",
        text,
    )
    lines: list[str] = []
    current = ""
    for token in tokens:
        if token == "\n":
            lines.append(current.rstrip())
            current = ""
            continue
        if token.isspace() and not current:
            continue
        candidate = current + token
        if current and pdf_text_width(candidate, font, size) > max_width:
            stripped_token = token.lstrip()
            if stripped_token and stripped_token[0] in no_line_start:
                # A closing mark may use the right inset, but never starts a line.
                current = candidate
                continue
            if current[-1] in no_line_end and len(current) > 1:
                lines.append(current[:-1].rstrip())
                current = current[-1] + stripped_token
            else:
                lines.append(current.rstrip())
                current = stripped_token
        else:
            current = candidate
        if current and pdf_text_width(current, font, size) > max_width:
            oversized = current
            current = ""
            parts = oversized.split()
            if len(parts) > 1:
                for part in parts:
                    candidate = f"{current} {part}" if current else part
                    if current and pdf_text_width(candidate, font, size) > max_width:
                        lines.append(current.rstrip())
                        current = part
                    else:
                        current = candidate
                    if pdf_text_width(current, font, size) <= max_width:
                        continue
                    unbreakable = current
                    current = ""
                    for char in unbreakable:
                        candidate = current + char
                        if current and pdf_text_width(candidate, font, size) > max_width:
                            lines.append(current.rstrip())
                            current = char.lstrip()
                        else:
                            current = candidate
            else:
                for char in oversized:
                    candidate = current + char
                    if current and pdf_text_width(candidate, font, size) > max_width:
                        lines.append(current.rstrip())
                        current = char.lstrip()
                    else:
                        current = candidate
    if current:
        lines.append(current.rstrip())
    return lines or [""]


def fit_pdf_lines(
    text: str,
    font: str,
    preferred_size: float,
    min_size: float,
    max_width: float,
    max_lines: int,
) -> tuple[float, list[str]]:
    """Fit complete text into a line budget; never truncate resume content."""
    size = preferred_size
    while size >= min_size - 0.001:
        lines = wrap_pdf_text(text, font, size, max_width)
        if len(lines) <= max_lines:
            return size, lines
        size = round(size - 0.1, 2)
    raise ValueError(
        f"PDF text cannot fit {max_lines} lines at readable size "
        f"({preferred_size:.1f}-{min_size:.1f}pt): {text}"
    )


def measure_bullet_content(items: list[str], width: float, available: float = 165) -> tuple[float, float, int]:
    """Return fitted font size, occupied height and wrapped line count."""
    size = 8.5
    leading = 10.9
    while size >= 7.6 - 0.001:
        line_count = sum(len(wrap_pdf_text(item, FONT_REGULAR, size, width - 34)) for item in items)
        required = line_count * leading + len(items) * 5
        if required <= available:
            return size, required, line_count
        size = round(size - 0.2, 2)
        leading = round(leading - 0.18, 2)
    raise ValueError("Bullet content cannot fit at the minimum readable size")


def choose_bullet_columns(
    content_width: float,
    gap: float,
    implementation: list[str],
    challenges: list[str],
) -> tuple[float, float]:
    """Balance the two engineering columns from their rendered line heights."""
    candidates: list[tuple[float, float, float]] = []
    available = 165.0
    for ratio in (0.58, 0.60, 0.62, 0.64, 0.65):
        implementation_width = content_width * ratio - gap / 2
        challenge_width = content_width - implementation_width - gap
        try:
            impl_size, impl_height, _ = measure_bullet_content(implementation, implementation_width, available)
            challenge_size, challenge_height, _ = measure_bullet_content(challenges, challenge_width, available)
        except ValueError:
            continue

        readability_penalty = (8.5 - min(impl_size, challenge_size)) * 220
        balance_penalty = abs(impl_height - challenge_height) * 0.7
        whitespace_penalty = max(0.0, available - min(impl_height, challenge_height)) * 0.25
        candidates.append(
            (readability_penalty + balance_penalty + whitespace_penalty, implementation_width, challenge_width)
        )

    if not candidates:
        raise ValueError("Engineering implementation and challenges do not fit any supported column ratio")
    _, implementation_width, challenge_width = min(candidates, key=lambda item: item[0])
    return implementation_width, challenge_width


class PdfPage:
    def __init__(self, c: canvas.Canvas, page_no: int, chapter: str):
        self.c = c
        self.page_no = page_no
        self.chapter = chapter
        self.width, self.height = A4
        self.rail = 42
        self.left = 60
        self.right = self.width - 34
        self.content_width = self.right - self.left
        self._draw_base()

    def _draw_base(self) -> None:
        c = self.c
        c.setFillColor(HexColor(f"#{COLD_WHITE}"))
        c.rect(0, 0, self.width, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.rect(0, 0, self.rail, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.rect(self.rail, 0, 3, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.rect(0, self.height - 112, self.rail, 36, fill=1, stroke=0)
        c.saveState()
        c.translate(16, 55)
        c.rotate(90)
        c.setFont(FONT_BOLD, 7.3)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawString(0, 0, "FU DAOPIN · JAVA BACKEND · ENGINEERING CASEBOOK")
        c.restoreState()
        c.setFont(FONT_BOLD, 12)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawCentredString(self.rail / 2, self.height - 101, f"{self.page_no:02d}")

    def title(self, title: str, subtitle: str) -> None:
        c = self.c
        c.setFont(FONT_BOLD, 7.8)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.drawString(self.left, self.height - 38, self.chapter.upper())
        c.setFont(FONT_BOLD, 19)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(self.left, self.height - 67, title)
        c.setFont(FONT_REGULAR, 8.6)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawString(self.left, self.height - 85, subtitle)
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.line(self.left, self.height - 98, self.right, self.height - 98)

    def section_label(self, x: float, y: float, label: str, index: str | None = None) -> None:
        c = self.c
        if index:
            c.setFont(FONT_BOLD, 6.8)
            c.setFillColor(HexColor(f"#{ORANGE}"))
            c.drawString(x, y, index)
            x += 18
        c.setFont(FONT_BOLD, 8.4)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.drawString(x, y, label)

    def wrapped(
        self,
        text: str,
        x: float,
        y: float,
        width: float,
        size: float = 8.8,
        leading: float = 12,
        color: str = INK,
        bold: bool = False,
        max_lines: int | None = None,
        min_size: float | None = None,
    ) -> float:
        font = FONT_BOLD if bold else FONT_REGULAR
        lines = wrap_pdf_text(text, font, size, width)
        if max_lines is not None and len(lines) > max_lines:
            fitted_size, lines = fit_pdf_lines(
                text,
                font,
                size,
                min_size if min_size is not None else max(6.8, size - 1.2),
                width,
                max_lines,
            )
            leading *= fitted_size / size
            size = fitted_size
        self.c.setFont(font, size)
        self.c.setFillColor(HexColor(f"#{color}"))
        for line in lines:
            self.c.drawString(x, y, line)
            y -= leading
        return y

    def info_box(self, x: float, y: float, width: float, height: float, label: str, text: str, accent: str = TEAL) -> None:
        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y + height - 3, width, 3, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 7.7)
        c.setFillColor(HexColor(f"#{accent}"))
        c.drawString(x + 10, y + height - 18, label)
        self.wrapped(
            text,
            x + 10,
            y + height - 34,
            width - 20,
            size=8.35,
            leading=11.2,
            max_lines=max(2, int((height - 31) / 11.2)),
            min_size=7.2,
        )

    def labeled_rows_box(
        self,
        x: float,
        y: float,
        width: float,
        height: float,
        rows: list[tuple[str, str, str]],
        preferred_size: float = 8.0,
        min_size: float = 7.4,
    ) -> None:
        label_width = 54
        text_x = x + label_width + 10
        text_width = width - label_width - 20
        size = preferred_size
        measured: list[list[str]] = []
        while size >= min_size - 0.001:
            leading = size + 1.8
            measured = [wrap_pdf_text(text, FONT_REGULAR, size, text_width) for _, text, _ in rows]
            required = 17 + sum(len(lines) * leading + 2 for lines in measured)
            if required <= height:
                break
            size = round(size - 0.1, 2)
        else:
            labels = " / ".join(label for label, _, _ in rows)
            raise ValueError(f"PDF page {self.page_no} labeled rows overflow: {labels}")

        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.rect(x, y + height - 3, width, 3, fill=1, stroke=0)

        leading = size + 1.8
        cursor = y + height - 13
        for index, ((label, _, accent), lines) in enumerate(zip(rows, measured)):
            c.setFont(FONT_BOLD, 7.5)
            c.setFillColor(HexColor(f"#{accent}"))
            c.drawString(x + 10, cursor, label)
            c.setFont(FONT_REGULAR, size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in lines:
                c.drawString(text_x, cursor, line)
                cursor -= leading
            cursor -= 2
            if index < len(rows) - 1:
                c.setStrokeColor(HexColor("#E2EAEC"))
                c.setLineWidth(0.35)
                c.line(x + 10, cursor + 1.5, x + width - 10, cursor + 1.5)

    def bullet_box(self, x: float, y: float, width: float, height: float, label: str, items: list[str], accent: str) -> None:
        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y + height - 4, width, 4, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 9)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(x + 12, y + height - 23, label)
        available = height - 48
        size = 8.5
        leading = 10.9
        line_count = 0
        required = 0.0
        while size >= 7.6 - 0.001:
            line_count = sum(len(wrap_pdf_text(item, FONT_REGULAR, size, width - 34)) for item in items)
            required = line_count * leading + len(items) * 5
            if required <= available:
                break
            size = round(size - 0.2, 2)
            leading = round(leading - 0.18, 2)
        else:
            raise ValueError(f"PDF page {self.page_no} bullet box overflow: {label}")

        # Keep dense engineering content readable while distributing modest breathing
        # room between bullets. The final baseline always retains an 8pt bottom inset.
        item_gap = min(10.0, max(5.0, (available - 8 - line_count * leading) / len(items)))
        cursor = y + height - 41
        for item in items:
            lines = wrap_pdf_text(item, FONT_REGULAR, size, width - 34)
            c.setFillColor(HexColor(f"#{accent}"))
            c.rect(x + 12, cursor + 2, 4, 4, fill=1, stroke=0)
            c.setFont(FONT_REGULAR, size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in lines:
                c.drawString(x + 23, cursor, line)
                cursor -= leading
            cursor -= item_gap
        if cursor < y + 7:
            raise ValueError(f"PDF page {self.page_no} bullet box overflow: {label}")

    def pipeline(self, y: float, flow: str) -> None:
        nodes = [item.strip() for item in flow.split("→")]
        gap = 8
        node_width = (self.content_width - gap * (len(nodes) - 1)) / len(nodes)
        for index, node in enumerate(nodes):
            x = self.left + index * (node_width + gap)
            self.c.setFillColor(HexColor("#E7F4F3" if index % 2 == 0 else "#EDF2F3"))
            self.c.setStrokeColor(HexColor(f"#{TEAL}"))
            self.c.roundRect(x, y, node_width, 31, 2, fill=1, stroke=1)
            node_size, lines = fit_pdf_lines(node, FONT_BOLD, 7.2, 6.4, node_width - 10, 2)
            self.c.setFont(FONT_BOLD, node_size)
            self.c.setFillColor(HexColor(f"#{NAVY}"))
            text_y = y + 18 + (len(lines) == 1) * -3
            for line in lines:
                self.c.drawCentredString(x + node_width / 2, text_y, line)
                text_y -= 9
            if index < len(nodes) - 1:
                arrow_x = x + node_width + 1
                self.c.setStrokeColor(HexColor(f"#{ORANGE}"))
                self.c.line(arrow_x, y + 15.5, arrow_x + gap - 2, y + 15.5)
                self.c.line(arrow_x + gap - 5, y + 18, arrow_x + gap - 2, y + 15.5)
                self.c.line(arrow_x + gap - 5, y + 13, arrow_x + gap - 2, y + 15.5)

    def footer(self) -> None:
        self.c.setFont(FONT_REGULAR, 7.2)
        self.c.setFillColor(HexColor(f"#{MUTED}"))
        self.c.drawRightString(self.right, 17, f"CASEBOOK  /  {self.page_no:02d} OF 10")


def draw_pdf_cover(c: canvas.Canvas, data: dict) -> None:
    w, h = A4
    c.setFillColor(HexColor(f"#{GRAPHITE}"))
    c.rect(0, 0, w, h, fill=1, stroke=0)
    for x in range(40, 561, 52):
        c.setStrokeColor(HexColor("#14272E"))
        c.setLineWidth(0.4)
        c.line(x, 0, x, h)
    for y in range(36, 806, 52):
        c.line(0, y, w, y)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.rect(0, 0, 15, h, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(15, h - 145, 9, 95, fill=1, stroke=0)

    # Industrial pipe/data schematic.
    c.setStrokeColor(HexColor(f"#{CYAN}"))
    c.setLineWidth(2.2)
    path = c.beginPath()
    path.moveTo(315, 690)
    path.lineTo(465, 690)
    path.lineTo(465, 575)
    path.lineTo(545, 575)
    c.drawPath(path, stroke=1, fill=0)
    for x, y, radius in ((315, 690, 7), (465, 690, 9), (465, 575, 6), (545, 575, 8)):
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.circle(x, y, radius, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{CYAN}"))
        c.circle(x, y, 2.2, fill=1, stroke=0)
    c.setStrokeColor(HexColor(f"#{ORANGE}"))
    c.setLineWidth(1.2)
    c.line(412, 736, 545, 736)
    c.line(545, 736, 545, 610)

    c.setFont(FONT_BOLD, 8.5)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawString(54, h - 67, "ENGINEERING CASEBOOK / 2026")
    c.setFont(FONT_BOLD, 35)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(54, h - 151, data["profile"]["name"])
    c.setFont(FONT_BOLD, 20)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(54, h - 189, data["profile"]["title"])
    c.setFont(FONT_REGULAR, 10)
    c.setFillColor(HexColor("#A8BAC0"))
    c.drawString(54, h - 215, data["profile"]["experience"])
    c.setStrokeColor(HexColor("#33515A"))
    c.line(54, h - 244, 541, h - 244)

    c.setFont(FONT_BOLD, 48)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawString(54, h - 337, "10w+")
    c.setFont(FONT_BOLD, 10)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(57, h - 359, "国内外水站 / GLOBAL WATER STATIONS")

    label_y = h - 435
    for index, item in enumerate(data["highlights"]):
        x = 54 + (index % 2) * 248
        y = label_y - (index // 2) * 57
        c.setFillColor(HexColor("#102027"))
        c.setStrokeColor(HexColor("#264750"))
        c.rect(x, y, 230, 43, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{TEAL}" if index != 3 else f"#{ORANGE}"))
        c.rect(x, y, 4, 43, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 8.2)
        c.setFillColor(HexColor(f"#{COLD_WHITE}"))
        highlight_size, lines = fit_pdf_lines(item, FONT_BOLD, 8.2, 7.4, 205, 2)
        c.setFont(FONT_BOLD, highlight_size)
        for line_index, line in enumerate(lines):
            c.drawString(x + 14, y + 25 - line_index * 11, line)

    c.setFont(FONT_REGULAR, 9.3)
    c.setFillColor(HexColor("#B8C8CC"))
    summary_lines = wrap_pdf_text(data["profile"]["summary"], FONT_REGULAR, 9.3, 487)
    sy = 200
    for line in summary_lines:
        c.drawString(54, sy, line)
        sy -= 14
    c.setStrokeColor(HexColor(f"#{TEAL}"))
    c.line(54, 145, 541, 145)
    c.setFont(FONT_BOLD, 9.2)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(54, 118, data["profile"]["phone"])
    c.drawString(190, 118, data["profile"]["email"])
    c.setFont(FONT_REGULAR, 7.3)
    c.setFillColor(HexColor("#6E858D"))
    c.drawRightString(541, 28, "01 / 10 · JAVA BACKEND · AIOT · GIS · AGENT")


def draw_pdf_overview(c: canvas.Canvas, data: dict) -> None:
    page = PdfPage(c, 2, "PROFILE / CAPABILITY MATRIX")
    page.title("核心简历", "高级 Java 开发工程师 · 微服务 / 智慧水务 / AIoT / GIS / Agent")
    page.info_box(page.left, 670, page.content_width, 58, "职业定位", data["profile"]["summary"], ORANGE)
    page.section_label(page.left, 649, "核心能力矩阵", "A1")
    box_gap = 10
    box_width = (page.content_width - box_gap) / 2
    for index, strength in enumerate(data["strengths"]):
        x = page.left + (index % 2) * (box_width + box_gap)
        y = 548 - (index // 2) * 91
        page.info_box(x, y, box_width, 78, strength["title"], strength["evidence"], TEAL if index < 3 else ORANGE)

    page.section_label(page.left, 443, "工作经历", "A2")
    timeline_x = page.left + 7
    c.setStrokeColor(HexColor(f"#{TEAL}"))
    c.setLineWidth(1.2)
    c.line(timeline_x, 170, timeline_x, 421)
    experience_top = 395
    experience_bottom = 215
    experience_step = (
        (experience_top - experience_bottom) / (len(data["experiences"]) - 1)
        if len(data["experiences"]) > 1
        else 0
    )
    for index, exp in enumerate(data["experiences"]):
        y = experience_top - index * experience_step
        c.setFillColor(HexColor(f"#{ORANGE}" if index == 0 else f"#{TEAL}"))
        c.circle(timeline_x, y + 9, 4.2, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 8.3)
        c.setFillColor(HexColor(f"#{ORANGE}" if index == 0 else f"#{MUTED}"))
        c.drawString(timeline_x + 17, y + 17, exp["period"])
        c.setFont(FONT_BOLD, 10.2)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(timeline_x + 122, y + 17, f"{exp['company']}  |  {exp['title']}")
        page.wrapped(exp["summary"], timeline_x + 17, y - 2, page.content_width - 24, size=8.25, leading=10.8, max_lines=3)

    stack: list[str] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)
    page.info_box(page.left, 73, page.content_width, 78, "技术域索引", " · ".join(stack[:24]), ORANGE)
    page.footer()


def draw_pdf_topic(c: canvas.Canvas, page_no: int, project: dict, topic: dict) -> None:
    chapter = "LITREE / ONE PROJECT · SIX ENGINEERING TOPICS" if project["id"] == "litree" else project["name"]
    page = PdfPage(c, page_no, chapter)
    page.title(topic["title"], f"{project['company']}  |  {project['name']}  |  {project['period']}")

    fact = project["summary"]
    if project["id"] == "litree":
        fact = "同一 Litree 项目六个工程专题展开 · 项目覆盖国内外 10w+ 水站"
    page.info_box(page.left, 681, page.content_width, 51, "项目定位 / 当前专题", fact, ORANGE)

    page.section_label(page.left, 661, "项目背景：业务场景 → 现有痛点 → 建设目标", "B1")
    gap = 8
    page.labeled_rows_box(
        page.left,
        513,
        page.content_width,
        133,
        [
            ("业务场景", project["businessContext"], TEAL),
            ("当前专题", topic["background"], CYAN),
            ("现有痛点", "；".join(item.rstrip("。；") for item in project["painPoints"]) + "。", ORANGE),
            ("建设目标", "；".join(item.rstrip("。；") for item in project["buildGoals"]) + "。", TEAL),
        ],
    )

    page.section_label(page.left, 500, "业务交付链路", "B2")
    page.pipeline(457, topic["flow"])

    page.section_label(page.left, 441, "职责与工程边界", "B3")
    page.labeled_rows_box(
        page.left,
        365,
        page.content_width,
        61,
        [
            ("本人角色", topic["role"], TEAL),
            ("工程边界", topic["engineeringBoundary"], ORANGE),
        ],
        preferred_size=8.2,
        min_size=7.4,
    )

    page.section_label(page.left, 351, "工程实现", "B4")
    implementation_width, challenge_width = choose_bullet_columns(
        page.content_width,
        gap,
        topic["implementation"],
        topic["challenges"],
    )
    page.bullet_box(page.left, 122, implementation_width, 213, "核心实现", topic["implementation"], TEAL)
    page.bullet_box(page.left + implementation_width + gap, 122, challenge_width, 213, "技术难点", topic["challenges"], ORANGE)

    c.setFillColor(HexColor(f"#{GRAPHITE}"))
    c.rect(page.left, 58, page.content_width, 51, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(page.left, 58, 5, 51, fill=1, stroke=0)
    c.setFont(FONT_BOLD, 7.8)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(page.left + 14, 90, "落地结果")
    page.wrapped(topic["outcome"], page.left + 14, 74, page.content_width - 28, size=8.45, leading=10.5, color=COLD_WHITE, max_lines=2)

    c.setFont(FONT_BOLD, 7.4)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.drawString(page.left, 43, "TECH MATRIX")
    stack_text = "  /  ".join(topic["stack"])
    page.wrapped(stack_text, page.left + 72, 43, page.content_width - 72, size=7.15, leading=8.8, color=MUTED, max_lines=2)
    page.footer()


def build_pdf(data: dict) -> None:
    register_pdf_fonts()
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4, pageCompression=1)
    c.setTitle("付道品 - 高级 Java 开发工程师")
    c.setAuthor("付道品")
    c.setSubject("高级 Java 开发工程师简历")
    draw_pdf_cover(c, data)
    c.showPage()
    draw_pdf_overview(c, data)

    pages: list[tuple[dict, dict]] = []
    for project in data["projects"]:
        for topic in project["topics"]:
            pages.append((project, topic))
    for page_no, (project, topic) in enumerate(pages, start=3):
        c.showPage()
        draw_pdf_topic(c, page_no, project, topic)
    c.save()


def main() -> None:
    data = load_data()
    build_docx(data)
    build_pdf(data)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
