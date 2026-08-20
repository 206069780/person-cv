from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]

OUTPUT_STEMS = {
    "zh": "付道品-高级Java开发工程师",
    "en": "Daopin-Fu-Senior-Java-Engineer",
}


def resume_bundle(lang: str) -> dict:
    stem = OUTPUT_STEMS[lang]
    return {
        "data": ROOT / "web" / "src" / "data" / f"resume-data.{lang}.json",
        "pdf": ROOT / f"{stem}.pdf",
        "docx": ROOT / f"{stem}.docx",
    }


FONT_REGULAR_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_REGULAR = "ResumeYaHei"
FONT_BOLD = "ResumeYaHeiBold"

GRAPHITE = "0B1418"
NAVY = "16323A"
TEAL = "0E8F86"
CYAN = "1AABB8"
ORANGE = "E24E24"
INK = "11181C"
MUTED = "4A5C64"
PALE = "EEF3F4"
RULE = "D5DEE1"
COLD_WHITE = "F4F6F7"
SURFACE = "FFFFFF"
RAIL_TEXT = "8FA3AB"
ON_DARK = "E8EEF0"
TEAL_SOFT = "E6F3F2"
ORANGE_SOFT = "FDECE6"
PAD = 8
GAP = 8
RADIUS = 3
TYPE_KICKER = 7.2
TYPE_TITLE = 18
TYPE_SECTION = 8.2
TYPE_BODY = 8.0
TYPE_CAPTION = 7.2
TYPE_CHIP = 6.4
LINE_RATIO = 1.58
WIDE_LINE_RATIO = 1.64
MIN_LINE_RATIO = 1.52
TOTAL_PAGES = 7

CHROME = {
    "zh": {
        "page_before": "第 ",
        "page_after": f" / {TOTAL_PAGES:02d} 页",
        "stack": "技术栈  ",
        "museum": "3D 展馆",
        "overview": "核心简历",
        "overview_sub": "面向普通投递的能力概览与工作经历",
        "strengths": "核心能力",
        "experience": "工作经历",
        "position": "项目定位",
        "business": "业务场景",
        "pain": "现有痛点",
        "goals": "建设目标",
        "role": "本人角色",
        "flow": "业务链路",
        "boundary": "工程边界",
        "modules": "核心业务模块",
        "module": "模块",
        "outcome": "落地结果",
        "practice": "核心技术实践",
        "implementation": "核心实现",
        "challenges": "技术难点",
        "module_outcome": "模块结果",
        "bg": "背景",
        "role_short": "角色",
        "flow_short": "链路",
        "boundary_short": "边界",
        "result": "结果",
        "cover_domains": "分布式微服务 / 智慧水务 / AIoT & GIS / Agent 工程化",
        "cover_water": "国内外水站 / GLOBAL WATER STATIONS",
        "cover_statement": "> ENGINEERING STATEMENT // 职业定位与工程准则",
        "cover_statement_body": "★ 严谨工程边界意识 · 注重系统可恢复性、数据一致性与生产可观测性 · 具备端到端落地交付经验",
        "cover_base": "BASE: 深圳 / 广州 · 全职",
        "cover_confidential": "CONFIDENTIAL / {name} 个人技术经历与工程案例集",
        "overview_kicker": "高级 Java 开发工程师 · 3D 展馆: {website} · 微服务 / AIoT / GIS / Agent",
        "overview_position": "职业定位",
        "overview_strengths": "核心能力矩阵",
        "overview_experience": "工作经历",
        "overview_stack": "技术域索引",
        "system_arch": "系统架构",
        "system_position": "项目定位 / 系统边界",
        "system_background": "系统背景：业务场景 → 现有痛点 → 建设目标",
        "heading_flow": "总体业务链路",
        "heading_duty": "职责与工程边界",
        "system_practice": "系统实践",
        "impl_challenges": "核心实现与技术难点",
        "module_impl": "核心实现",
        "pdf_subject": "高级 Java 开发工程师简历",
        "item_join": "；",
        "item_stop": "。",
        "item_strip": "。；",
    },
    "en": {
        "page_before": "Page ",
        "page_after": f" / {TOTAL_PAGES:02d}",
        "stack": "Stack  ",
        "museum": "3D museum",
        "overview": "Core resume",
        "overview_sub": "Competency overview and work history for general applications",
        "strengths": "Core skills",
        "experience": "Work experience",
        "position": "Positioning",
        "business": "Business context",
        "pain": "Pain points",
        "goals": "Build goals",
        "role": "Role",
        "flow": "Flow",
        "boundary": "Engineering boundary",
        "modules": "Core modules",
        "module": "Module",
        "outcome": "Outcome",
        "practice": "Core technical practice",
        "implementation": "Implementation",
        "challenges": "Challenges",
        "module_outcome": "Module outcome",
        "bg": "Background",
        "role_short": "Role",
        "flow_short": "Flow",
        "boundary_short": "Boundary",
        "result": "Result",
        "cover_domains": "Distributed microservices / smart water / AIoT & GIS / Agent engineering",
        "cover_water": "Water stations worldwide / GLOBAL WATER STATIONS",
        "cover_statement": "> ENGINEERING STATEMENT // positioning and engineering bar",
        "cover_statement_body": "★ Strict ownership boundaries · recoverability, consistency, and production observability · end-to-end delivery",
        "cover_base": "BASE: Shenzhen / Guangzhou · full-time",
        "cover_confidential": "CONFIDENTIAL / {name} engineering casebook",
        "overview_kicker": "Senior Java Engineer · 3D museum: {website} · microservices / AIoT / GIS / Agent",
        "overview_position": "Positioning",
        "overview_strengths": "Core competency matrix",
        "overview_experience": "Work experience",
        "overview_stack": "Technology index",
        "system_arch": "System architecture",
        "system_position": "Positioning / system boundary",
        "system_background": "Background: context → pain points → goals",
        "heading_flow": "End-to-end flow",
        "heading_duty": "Role and engineering boundary",
        "system_practice": "System practice",
        "impl_challenges": "Implementation and challenges",
        "module_impl": "Implementation",
        "pdf_subject": "Senior Java Engineer resume",
        "item_join": "; ",
        "item_strip": ".;",
        "item_stop": ".",
    },
}


def load_data(lang: str) -> dict:
    return json.loads(resume_bundle(lang)["data"].read_text(encoding="utf-8"))


def set_run_font(run, size: float, bold: bool = False, color: str = INK) -> None:
    name = "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeatable_cell_margins(cell, top=70, start=100, bottom=70, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph, lang: str) -> None:
    chrome = CHROME[lang]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(chrome["page_before"])
    set_run_font(run, 8, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    end = paragraph.add_run(chrome["page_after"])
    set_run_font(end, 8, color=MUTED)


def configure_docx(doc: Document, data: dict, lang: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.5

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(7)

    for style_name, size, before, after, color in (
        ("Heading 1", 17, 0, 8, NAVY),
        ("Heading 2", 11.5, 7, 4, TEAL),
        ("Heading 3", 9.5, 5, 2, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Cm(0.56)
        style.paragraph_format.first_line_indent = Cm(-0.28)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    if "Resume Lead" not in [s.name for s in styles]:
        lead = styles.add_style("Resume Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Resume Lead"]
    lead.font.name = "Microsoft YaHei"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lead.font.size = Pt(10.5)
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.58

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{data['profile']['name']}  |  {data['profile']['title']}")
    set_run_font(r, 8, bold=True, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0], lang)


def add_docx_title(doc: Document, kicker: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(kicker.upper())
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 17, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(subtitle)
        set_run_font(r, 9, color=MUTED)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.58
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{label}  ")
    set_run_font(r, 8.8, bold=True, color=TEAL)
    r = p.add_run(text)
    set_run_font(r, 8.7, color=INK)


def add_docx_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run_font(r, 8.5, color=INK)


def add_stack(doc: Document, stack: list[str], lang: str) -> None:
    chrome = CHROME[lang]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(chrome["stack"])
    set_run_font(r, 8.5, bold=True, color=ORANGE)
    r = p.add_run(" · ".join(stack))
    set_run_font(r, 8.2, color=MUTED)


def add_docx_cover(doc: Document, data: dict, lang: str) -> None:
    chrome = CHROME[lang]
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("JAVA BACKEND · AIOT · AGENT ENGINEERING")
    set_run_font(r, 9.5, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(data["profile"]["name"])
    set_run_font(r, 31, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(data["profile"]["title"])
    set_run_font(r, 17, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(data["profile"]["experience"])
    set_run_font(r, 10.5, color=MUTED)

    cols = 2
    rows = (len(data["highlights"]) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    col_width = Cm(17.9 / cols)
    for idx, value in enumerate(data["highlights"]):
        r_idx = idx // cols
        c_idx = idx % cols
        cell = table.cell(r_idx, c_idx)
        cell.width = col_width
        set_repeatable_cell_margins(cell, 80, 80, 80, 80)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), PALE)
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 8.5, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data["profile"]["summary"])
    set_run_font(r, 10.2, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    website = data["profile"].get("website", "http://cv.bookfree.online/")
    r = p.add_run(f"{data['profile']['phone']}  |  {data['profile']['email']}  |  {chrome['museum']}: {website}")
    set_run_font(r, 9.5, bold=True, color=TEAL)


def add_docx_overview(doc: Document, data: dict, lang: str) -> None:
    chrome = CHROME[lang]
    add_docx_title(doc, f"02 / {TOTAL_PAGES:02d}", chrome["overview"], chrome["overview_sub"])
    p = doc.add_paragraph(style="Resume Lead")
    p.add_run(data["profile"]["summary"])

    doc.add_paragraph(chrome["strengths"], style="Heading 2")
    for strength in data["strengths"]:
        add_labeled_paragraph(doc, strength["title"], strength["evidence"])

    doc.add_paragraph(chrome["experience"], style="Heading 2")
    for exp in data["experiences"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{exp['period']}  |  {exp['company']}  |  {exp['title']}")
        set_run_font(r, 9.1, bold=True, color=NAVY)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(exp["summary"])
        set_run_font(r, 8.4)
        add_docx_bullets(doc, exp["achievements"])

    stack = []
    for project in data["projects"]:
        for item in project.get("stack", []):
            if item not in stack:
                stack.append(item)
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)
    add_stack(doc, stack[:24], lang)


def join_items(items: list[str], lang: str) -> str:
    chrome = CHROME[lang]
    return chrome["item_join"].join(item.rstrip(chrome["item_strip"]) for item in items) + chrome["item_stop"]


def add_docx_system_overview(doc: Document, page_no: int, project: dict, lang: str) -> None:
    chrome = CHROME[lang]
    add_docx_title(
        doc,
        f"{page_no:02d} / {TOTAL_PAGES:02d} · {project['company']}",
        project["name"],
        project["period"],
    )
    add_labeled_paragraph(doc, chrome["position"], project["summary"])
    add_labeled_paragraph(doc, chrome["business"], project["businessContext"])
    add_labeled_paragraph(doc, chrome["pain"], join_items(project["painPoints"], lang))
    add_labeled_paragraph(doc, chrome["goals"], join_items(project["buildGoals"], lang))
    add_labeled_paragraph(doc, chrome["role"], project["role"])
    add_labeled_paragraph(doc, chrome["flow"], project["flow"])
    add_labeled_paragraph(doc, chrome["boundary"], project["engineeringBoundary"])
    doc.add_paragraph(chrome["modules"], style="Heading 2")
    for index, topic in enumerate(project["topics"], start=1):
        add_labeled_paragraph(doc, f"{chrome['module']} {index:02d}  {topic['title']}", topic["background"])
    add_labeled_paragraph(doc, chrome["outcome"], project["outcome"])
    add_stack(doc, project["stack"], lang)


def add_docx_system_modules(doc: Document, page_no: int, project: dict, lang: str) -> None:
    chrome = CHROME[lang]
    add_docx_title(
        doc,
        f"{page_no:02d} / {TOTAL_PAGES:02d} · {project['company']}",
        f"{project['name']} · {chrome['practice']}",
        project["period"],
    )
    for index, topic in enumerate(project["topics"], start=1):
        doc.add_paragraph(f"{chrome['module']} {index:02d}  {topic['title']}", style="Heading 2")
        add_labeled_paragraph(doc, chrome["role"], topic["role"])
        add_labeled_paragraph(doc, chrome["flow"], topic["flow"])
        doc.add_paragraph(chrome["implementation"], style="Heading 3")
        add_docx_bullets(doc, topic["implementation"])
        doc.add_paragraph(chrome["challenges"], style="Heading 3")
        add_docx_bullets(doc, topic["challenges"])
        add_labeled_paragraph(doc, chrome["outcome"], topic["outcome"])
    add_stack(doc, project["stack"], lang)


def add_docx_system_onepager(doc: Document, page_no: int, project: dict, lang: str) -> None:
    chrome = CHROME[lang]
    add_docx_title(
        doc,
        f"{page_no:02d} / {TOTAL_PAGES:02d} · {project['company']}",
        project["name"],
        project["period"],
    )
    add_labeled_paragraph(doc, chrome["position"], project["summary"])
    add_labeled_paragraph(doc, chrome["business"], project["businessContext"])
    add_labeled_paragraph(doc, chrome["pain"], join_items(project["painPoints"], lang))
    add_labeled_paragraph(doc, chrome["goals"], join_items(project["buildGoals"], lang))
    add_labeled_paragraph(doc, chrome["role"], project["role"])
    add_labeled_paragraph(doc, chrome["flow"], project["flow"])
    add_labeled_paragraph(doc, chrome["boundary"], project["engineeringBoundary"])
    for index, topic in enumerate(project["topics"], start=1):
        doc.add_paragraph(f"{chrome['module']} {index:02d}  {topic['title']}", style="Heading 2")
        add_docx_bullets(doc, topic["implementation"])
        add_labeled_paragraph(doc, chrome["challenges"], chrome["item_join"].join(topic["challenges"]))
        add_labeled_paragraph(doc, chrome["module_outcome"], topic["outcome"])
    add_labeled_paragraph(doc, chrome["outcome"], project["outcome"])
    add_stack(doc, project["stack"], lang)


def build_docx(data: dict, lang: str, out_path: Path) -> None:
    doc = Document()
    configure_docx(doc, data, lang)
    add_docx_cover(doc, data, lang)
    doc.add_page_break()
    add_docx_overview(doc, data, lang)

    projects = data["projects"]
    if [item["id"] for item in projects] != ["litree", "oa", "welink", "senge"]:
        raise ValueError("Expected systems: litree, oa, welink, senge")
    if [item["pageSpan"] for item in projects] != [2, 1, 1, 1]:
        raise ValueError("Expected page spans: 2, 1, 1, 1")

    doc.add_page_break()
    add_docx_system_overview(doc, 3, projects[0], lang)
    doc.add_page_break()
    add_docx_system_modules(doc, 4, projects[0], lang)
    for page_no, project in enumerate(projects[1:], start=5):
        doc.add_page_break()
        add_docx_system_onepager(doc, page_no, project, lang)
    doc.save(out_path)


def register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont(FONT_REGULAR, str(FONT_REGULAR_PATH)))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, str(FONT_BOLD_PATH)))


def pdf_text_width(text: str, font: str, size: float) -> float:
    return pdfmetrics.stringWidth(text, font, size)


def line_leading(size: float, wide: bool = False) -> float:
    ratio = WIDE_LINE_RATIO if wide else LINE_RATIO
    return round(size * max(MIN_LINE_RATIO, ratio), 2)


def paragraph_gap(size: float) -> float:
    return round(max(6.0, size * 0.62), 2)


def wrap_pdf_text(text: str, font: str, size: float, max_width: float) -> list[str]:
    no_line_start = set("，。！？；：、）》】〕」』”’％%)]},.;:!?/\\")
    no_line_end = set("（《【〔「『“‘([{")
    tokens = re.findall(
        r"[A-Za-z0-9_+#’'-]+|[ \t]+|\n|/|[^\s\w]|\w",
        text,
    )
    lines: list[str] = []
    current = ""
    for token in tokens:
        if token == "\n":
            lines.append(current.rstrip())
            current = ""
            continue
        if token.isspace() and not current:
            continue
        candidate = current + token
        if current and pdf_text_width(candidate, font, size) > max_width:
            stripped_token = token.lstrip()
            if stripped_token and stripped_token[0] in no_line_start:
                # 避头符号（如 / 、逗号、句号）优先依附在上一行末尾
                current = candidate
                continue
            if current[-1] in no_line_end and len(current) > 1:
                lines.append(current[:-1].rstrip())
                current = current[-1] + stripped_token
            else:
                lines.append(current.rstrip())
                current = stripped_token
        else:
            current = candidate
        if current and pdf_text_width(current, font, size) > max_width:
            oversized = current
            current = ""
            parts = oversized.split()
            if len(parts) > 1:
                for part in parts:
                    candidate = f"{current} {part}" if current else part
                    if current and pdf_text_width(candidate, font, size) > max_width:
                        lines.append(current.rstrip())
                        current = part
                    else:
                        current = candidate
                    if pdf_text_width(current, font, size) <= max_width:
                        continue
                    unbreakable = current
                    current = ""
                    for char in unbreakable:
                        candidate = current + char
                        if current and pdf_text_width(candidate, font, size) > max_width:
                            lines.append(current.rstrip())
                            current = char.lstrip()
                        else:
                            current = candidate
            else:
                for char in oversized:
                    candidate = current + char
                    if current and pdf_text_width(candidate, font, size) > max_width:
                        lines.append(current.rstrip())
                        current = char.lstrip()
                    else:
                        current = candidate
    if current:
        lines.append(current.rstrip())
    return lines or [""]


def fit_pdf_lines(
    text: str,
    font: str,
    preferred_size: float,
    min_size: float,
    max_width: float,
    max_lines: int,
) -> tuple[float, list[str]]:
    """Fit complete text into a line budget; never truncate resume content."""
    size = preferred_size
    while size >= min_size - 0.001:
        lines = wrap_pdf_text(text, font, size, max_width)
        if len(lines) <= max_lines:
            return size, lines
        size = round(size - 0.1, 2)
    raise ValueError(
        f"PDF text cannot fit {max_lines} lines at readable size "
        f"({preferred_size:.1f}-{min_size:.1f}pt): {text}"
    )


def measure_bullet_content(items: list[str], width: float, available: float = 165) -> tuple[float, float, int]:
    """Return fitted font size, occupied height and wrapped line count."""
    size = 8.5
    while size >= 7.2 - 0.001:
        leading = line_leading(size)
        gap = paragraph_gap(size)
        line_count = sum(len(wrap_pdf_text(item, FONT_REGULAR, size, width - 28)) for item in items)
        required = line_count * leading + len(items) * gap
        if required <= available:
            return size, required, line_count
        size = round(size - 0.2, 2)
    raise ValueError("Bullet content cannot fit at the minimum readable size")


def choose_bullet_columns(
    content_width: float,
    gap: float,
    implementation: list[str],
    challenges: list[str],
) -> tuple[float, float]:
    """Balance the two engineering columns from their rendered line heights."""
    candidates: list[tuple[float, float, float]] = []
    available = 165.0
    for ratio in (0.58, 0.60, 0.62, 0.64, 0.65):
        implementation_width = content_width * ratio - gap / 2
        challenge_width = content_width - implementation_width - gap
        try:
            impl_size, impl_height, _ = measure_bullet_content(implementation, implementation_width, available)
            challenge_size, challenge_height, _ = measure_bullet_content(challenges, challenge_width, available)
        except ValueError:
            continue

        readability_penalty = (8.5 - min(impl_size, challenge_size)) * 220
        balance_penalty = abs(impl_height - challenge_height) * 0.7
        whitespace_penalty = max(0.0, available - min(impl_height, challenge_height)) * 0.25
        candidates.append(
            (readability_penalty + balance_penalty + whitespace_penalty, implementation_width, challenge_width)
        )

    if not candidates:
        raise ValueError("Engineering implementation and challenges do not fit any supported column ratio")
    _, implementation_width, challenge_width = min(candidates, key=lambda item: item[0])
    return implementation_width, challenge_width


def draw_round_card(c: canvas.Canvas, x: float, y: float, width: float, height: float, accent: str | None = None) -> None:
    c.setFillColor(HexColor(f"#{SURFACE}"))
    c.setStrokeColor(HexColor(f"#{RULE}"))
    c.setLineWidth(0.7)
    c.roundRect(x, y, width, height, RADIUS, fill=1, stroke=1)
    if accent:
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y + 1, 3.5, height - 2, fill=1, stroke=0)


def draw_index_badge(c: canvas.Canvas, x: float, y: float, index: str) -> float:
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.roundRect(x, y - 2.5, 16, 12, 2, fill=1, stroke=0)
    c.setFont(FONT_BOLD, 6.6)
    c.setFillColor(HexColor("#FFFFFF"))
    c.drawCentredString(x + 8, y + 0.6, index)
    return x + 22


class PdfPage:
    def __init__(self, c: canvas.Canvas, page_no: int, chapter: str, name: str):
        self.c = c
        self.page_no = page_no
        self.chapter = chapter
        self.name = name
        self.width, self.height = A4
        self.rail = 42
        self.left = 60
        self.right = self.width - 34
        self.content_width = self.right - self.left
        self._draw_base()

    def _draw_base(self) -> None:
        c = self.c
        c.setFillColor(HexColor(f"#{COLD_WHITE}"))
        c.rect(0, 0, self.width, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.rect(0, 0, self.rail, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.rect(self.rail, 0, 3, self.height, fill=1, stroke=0)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.rect(0, self.height - 108, self.rail, 32, fill=1, stroke=0)
        c.saveState()
        c.translate(15, 52)
        c.rotate(90)
        c.setFont(FONT_BOLD, 7.0)
        c.setFillColor(HexColor(f"#{RAIL_TEXT}"))
        rail_name = " ".join(self.name.upper().split())
        c.drawString(0, 0, f"{rail_name}  ·  JAVA BACKEND  ·  ENGINEERING CASEBOOK")
        c.restoreState()
        c.setFont(FONT_BOLD, 12)
        c.setFillColor(HexColor("#FFFFFF"))
        c.drawCentredString(self.rail / 2, self.height - 97, f"{self.page_no:02d}")

    def title(self, title: str, subtitle: str) -> None:
        c = self.c
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.rect(self.left, self.height - 38, 16, 2.8, fill=1, stroke=0)
        c.setFont(FONT_BOLD, TYPE_KICKER)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawString(self.left + 22, self.height - 38, self.chapter.upper())
        c.setFont(FONT_BOLD, TYPE_TITLE)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(self.left, self.height - 62, title)
        c.setFont(FONT_REGULAR, 8.2)
        c.setFillColor(HexColor(f"#{MUTED}"))
        c.drawString(self.left, self.height - 78, subtitle)
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.line(self.left, self.height - 88, self.right, self.height - 88)

    def section_label(self, x: float, y: float, label: str, index: str | None = None) -> None:
        c = self.c
        if index:
            x = draw_index_badge(c, x, y, index)
        c.setFont(FONT_BOLD, TYPE_SECTION)
        c.setFillColor(HexColor(f"#{NAVY}"))
        c.drawString(x, y, label)

    def wrapped(
        self,
        text: str,
        x: float,
        y: float,
        width: float,
        size: float = 8.8,
        leading: float | None = None,
        color: str = INK,
        bold: bool = False,
        max_lines: int | None = None,
        min_size: float | None = None,
    ) -> float:
        font = FONT_BOLD if bold else FONT_REGULAR
        if leading is None:
            leading = line_leading(size, wide=True)
        lines = wrap_pdf_text(text, font, size, width)
        if max_lines is not None and len(lines) > max_lines:
            fitted_size, lines = fit_pdf_lines(
                text,
                font,
                size,
                min_size if min_size is not None else max(6.8, size - 1.2),
                width,
                max_lines,
            )
            leading = line_leading(fitted_size, wide=True)
            size = fitted_size
        self.c.setFont(font, size)
        self.c.setFillColor(HexColor(f"#{color}"))
        for line in lines:
            self.c.drawString(x, y, line)
            y -= leading
        return y

    def info_box(self, x: float, y: float, width: float, height: float, label: str, text: str, accent: str = TEAL) -> None:
        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.roundRect(x, y, width, height, RADIUS, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x + 1, y + height - 3, width - 2, 3, fill=1, stroke=0)
        pad_x = 12
        text_width = width - pad_x * 2
        available_h = height - 3.0
        label_gap = 4.5

        # 动态自适应标题字号与折行，杜绝长英文标题溢出卡片右侧
        label_size = 7.5
        while label_size >= 6.2 and pdf_text_width(label, FONT_BOLD, label_size) > text_width:
            label_size = round(label_size - 0.2, 2)
        label_lines = wrap_pdf_text(label, FONT_BOLD, label_size, text_width)
        label_lead = round(label_size * 1.3, 2)
        label_block_h = (len(label_lines) - 1) * label_lead + label_size * 0.72

        # 严格自适应拟合：同时调节字号与行间距，确保内容绝对不溢出且留有上下安全边距
        preferred_size = 8.1 if width > 300 else 7.8
        min_size = 6.0
        ratio = WIDE_LINE_RATIO
        min_ratio = 1.36
        fitted = False
        lines: list[str] = []
        lead = round(preferred_size * ratio, 2)
        total_content_h = 0.0

        max_allowed_lines = 2 if height <= 60 and width > 300 else 4

        while ratio >= min_ratio - 0.001:
            size = preferred_size
            while size >= min_size - 0.001:
                lead = round(size * ratio, 2)
                lines = wrap_pdf_text(text, FONT_REGULAR, size, text_width)
                if len(lines) <= max_allowed_lines:
                    text_block_h = (len(lines) - 1) * lead + size * 0.72
                    total_content_h = label_block_h + label_gap + text_block_h
                    if total_content_h + 5.0 <= available_h:
                        fitted = True
                        break
                size = round(size - 0.1, 2)
            if fitted:
                break
            ratio = round(ratio - 0.04, 2)

        if not fitted:
            size = min_size
            lead = round(size * min_ratio, 2)
            lines = wrap_pdf_text(text, FONT_REGULAR, size, text_width)
            text_block_h = (len(lines) - 1) * lead + size * 0.72
            total_content_h = label_block_h + label_gap + text_block_h

        # 垂直绝对居中计算
        pad_top = max(2.5, (available_h - total_content_h) / 2)

        cur_label_y = y + height - 3.0 - pad_top - label_size * 0.72
        c.setFont(FONT_BOLD, label_size)
        c.setFillColor(HexColor(f"#{NAVY}"))
        for l_line in label_lines:
            c.drawString(x + pad_x, cur_label_y, l_line)
            cur_label_y -= label_lead

        text_y = cur_label_y + label_lead - label_gap - size * 0.72
        c.setFont(FONT_REGULAR, size)
        c.setFillColor(HexColor(f"#{INK}"))
        for line in lines:
            c.drawString(x + pad_x, text_y, line)
            text_y -= lead

    def experience_card(self, x: float, y: float, width: float, height: float, exp: dict, accent: str) -> None:
        c = self.c
        draw_round_card(c, x, y, width, height, accent)
        pad_x = 12
        header_y = y + height - 14
        period_str = exp["period"]
        c.setFont(FONT_BOLD, 7.8)
        c.setFillColor(HexColor(f"#{ORANGE}" if accent == ORANGE else f"#{TEAL}"))
        c.drawString(x + pad_x, header_y, period_str)

        period_w = pdf_text_width(period_str, FONT_BOLD, 7.8)
        comp_x = x + pad_x + max(92.0, period_w + 12.0)
        comp_w = (x + width - pad_x) - comp_x
        comp_title_str = f"{exp['company']}  |  {exp['title']}"
        title_font_size = 9.8
        while title_font_size >= 7.8 and pdf_text_width(comp_title_str, FONT_BOLD, title_font_size) > comp_w:
            title_font_size = round(title_font_size - 0.2, 2)

        c.setFont(FONT_BOLD, title_font_size)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(comp_x, header_y, comp_title_str)

        size = 8.0
        lead = round(size * 1.45, 2)
        cursor = header_y - 13
        cursor = self.wrapped(
            exp["summary"],
            x + pad_x,
            cursor,
            width - pad_x * 2,
            size=size,
            leading=lead,
            max_lines=2,
            min_size=7.2,
        )
        floor = y + 8
        achievements = exp.get("achievements", [])
        wrap_w = width - pad_x * 2 - 14
        bullet_size = 7.7
        bullet_lead = line_leading(bullet_size)
        item_gap = 5.0
        packed: list[list[str]] = []

        def pack(size_n: float, lead_n: float, gap_n: float) -> float:
            nonlocal packed
            packed = [wrap_pdf_text(item, FONT_REGULAR, size_n, wrap_w) for item in achievements]
            return sum(len(lines) * lead_n for lines in packed) + gap_n * max(0, len(packed) - 1)

        intro = 5.0
        available = cursor - intro - floor
        used = pack(bullet_size, bullet_lead, item_gap)
        while used > available and bullet_size > 6.4:
            bullet_size = round(bullet_size - 0.1, 2)
            bullet_lead = line_leading(bullet_size)
            used = pack(bullet_size, bullet_lead, item_gap)
        while used > available and item_gap > 2.0:
            item_gap = round(item_gap - 0.5, 2)
            used = pack(bullet_size, bullet_lead, item_gap)

        leftover = max(0.0, available - used)
        n_items = max(1, len(packed))
        cursor -= intro + leftover / (n_items + 1)
        extra_between = leftover / (n_items + 1)

        for index, lines in enumerate(packed):
            c.setFillColor(HexColor(f"#{accent}"))
            c.circle(x + pad_x + 3.2, cursor + bullet_size * 0.32, 1.5, fill=1, stroke=0)
            c.setFont(FONT_REGULAR, bullet_size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in lines:
                c.drawString(x + pad_x + 12, cursor, line)
                cursor -= bullet_lead
            if index < n_items - 1:
                cursor -= item_gap + extra_between

    def labeled_rows_box(
        self,
        x: float,
        y: float,
        width: float,
        height: float,
        rows: list[tuple[str, str, str]],
        preferred_size: float = 8.0,
        min_size: float = 6.4,
    ) -> None:
        label_size = 7.5
        # 动态测量当前所有 label 的最大实际宽度，绝不硬编码 54pt，彻底杜绝中英文重合与挤压
        max_label_w = max((pdf_text_width(label, FONT_BOLD, label_size) for label, _, _ in rows), default=40.0)
        label_col_w = max(52.0, max_label_w + 5.0)
        text_x = x + 10 + label_col_w + 5.0
        text_width = width - (10 + label_col_w + 5.0) - 10

        n_rows = len(rows)
        size = preferred_size
        ratio = WIDE_LINE_RATIO
        measured: list[list[str]] = []
        leading = line_leading(size, wide=True)
        fitted = False
        available_h = height - 3.0

        while ratio >= MIN_LINE_RATIO - 0.001:
            size = preferred_size
            while size >= min_size - 0.001:
                leading = round(size * ratio, 2)
                measured = [wrap_pdf_text(text, FONT_REGULAR, size, text_width) for _, text, _ in rows]
                h_list = [(len(lines) - 1) * leading + size * 0.72 for lines in measured]
                min_row_pad = 5.0 * n_rows
                if sum(h_list) + min_row_pad <= available_h:
                    fitted = True
                    break
                size = round(size - 0.1, 2)
            if fitted:
                break
            ratio = round(ratio - 0.02, 2)

        if not fitted:
            size = min_size
            leading = round(size * MIN_LINE_RATIO, 2)
            measured = [wrap_pdf_text(text, FONT_REGULAR, size, text_width) for _, text, _ in rows]
            h_list = [(len(lines) - 1) * leading + size * 0.72 for lines in measured]
            if sum(h_list) + 3.0 * n_rows > available_h:
                labels = " / ".join(label for label, _, _ in rows)
                raise ValueError(
                    f"PDF page {self.page_no} labeled rows overflow: {labels} in {height}"
                )

        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{TEAL}"))
        c.rect(x, y + height - 3, width, 3, fill=1, stroke=0)

        # 核心算法：每个 row 分配独立的垂直槽位，文字在各槽位中绝对垂直居中
        h_list = [(len(lines) - 1) * leading + size * 0.72 for lines in measured]
        total_text_h = sum(h_list)
        extra_total = available_h - total_text_h
        slot_extra = extra_total / n_rows

        current_top = y + height - 3.0
        for index, ((label, _, accent), lines, text_h) in enumerate(zip(rows, measured, h_list)):
            slot_h = text_h + slot_extra
            row_pad = slot_extra / 2.0
            baseline_1 = current_top - row_pad - size * 0.72

            c.setFont(FONT_BOLD, label_size)
            c.setFillColor(HexColor(f"#{accent}"))
            c.drawString(x + 10, baseline_1, label)

            c.setFont(FONT_REGULAR, size)
            c.setFillColor(HexColor(f"#{INK}"))
            line_y = baseline_1
            for line in lines:
                c.drawString(text_x, line_y, line)
                line_y -= leading

            current_top -= slot_h
            if index < n_rows - 1:
                c.setStrokeColor(HexColor("#E2EAEC"))
                c.setLineWidth(0.35)
                c.line(x + 10, current_top, x + width - 10, current_top)

    def bullet_box(self, x: float, y: float, width: float, height: float, label: str, items: list[str], accent: str) -> None:
        c = self.c
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y + height - 4, width, 4, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 9)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.drawString(x + 12, y + height - 23, label)
        available = height - 48
        size = 8.5
        line_count = 0
        required = 0.0
        leading = line_leading(size)
        gap = paragraph_gap(size)
        wrap_w = width - 28
        while size >= 7.2 - 0.001:
            leading = line_leading(size)
            gap = paragraph_gap(size)
            line_count = sum(len(wrap_pdf_text(item, FONT_REGULAR, size, wrap_w)) for item in items)
            required = line_count * leading + len(items) * gap
            if required <= available:
                break
            size = round(size - 0.2, 2)
        else:
            raise ValueError(f"PDF page {self.page_no} bullet box overflow: {label}")

        item_gap = min(8.5, max(gap, (available - 8 - line_count * leading) / len(items)))
        cursor = y + height - 41
        text_x = x + 24
        for item in items:
            lines = wrap_pdf_text(item, FONT_REGULAR, size, wrap_w)
            c.setFillColor(HexColor(f"#{accent}"))
            c.circle(x + 16, cursor + 1.8, 1.7, fill=1, stroke=0)
            c.setFont(FONT_REGULAR, size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in lines:
                c.drawString(text_x, cursor, line)
                cursor -= leading
            cursor -= item_gap
        if cursor < y + 7:
            raise ValueError(f"PDF page {self.page_no} bullet box overflow: {label}")

    def draw_stack_chips(self, x: float, y: float, width: float, height: float, tags: list[str]) -> None:
        if height < 16 or not tags:
            return
        c = self.c
        chip_y = y + height - 13
        chip_x = x
        chip_h = 11
        for tag in tags:
            chip_w = pdf_text_width(tag, FONT_REGULAR, 6.1) + 8
            if chip_x + chip_w > x + width:
                chip_x = x
                chip_y -= chip_h + 3
            if chip_y < y + 1:
                break
            c.setFillColor(HexColor("#E7F4F3"))
            c.setStrokeColor(HexColor(f"#{TEAL}"))
            c.setLineWidth(0.35)
            c.roundRect(chip_x, chip_y - 2, chip_w, chip_h, 1.5, fill=1, stroke=1)
            c.setFont(FONT_REGULAR, 6.1)
            c.setFillColor(HexColor(f"#{NAVY}"))
            c.drawString(chip_x + 4, chip_y + 1.1, tag)
            chip_x += chip_w + 3.5

    def draw_card_fill(self, x: float, y: float, width: float, height: float, background: str, tags: list[str], lang: str) -> None:
        if height < 22:
            return
        c = self.c
        chrome = CHROME[lang]
        c.setFillColor(HexColor("#E8F3F2"))
        c.setStrokeColor(HexColor(f"#{TEAL}"))
        c.setLineWidth(0.55)
        c.roundRect(x - 2, y, width + 4, height, 2, fill=1, stroke=1)
        cursor = y + height - 11
        chip_reserve = 28 if tags and height >= 40 else 0
        text_h = height - chip_reserve - 6
        if background and text_h >= 14:
            c.setFont(FONT_BOLD, 6.3)
            c.setFillColor(HexColor(f"#{TEAL}"))
            c.drawString(x, cursor, chrome["bg"])
            cursor -= 11
            size = 6.7
            lines = wrap_pdf_text(background, FONT_REGULAR, size, width - 2)
            text_budget = max(8.0, cursor - (y + chip_reserve + 4))
            max_lines = max(1, min(len(lines), int(text_budget / (size + 0.8))))
            lead = line_leading(size)
            lead = min(max(size * 1.45, lead), size + 3.8)
            c.setFont(FONT_REGULAR, size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in lines[:max_lines]:
                if cursor < y + chip_reserve + 6:
                    break
                c.drawString(x, cursor, line)
                cursor -= lead
        if tags:
            self.draw_stack_chips(x, y + 3, width, max(16.0, cursor - y - 4), tags)

    def module_card(
        self,
        x: float,
        y: float,
        width: float,
        height: float,
        index: int,
        topic: dict,
        accent: str,
        lang: str,
    ) -> None:
        c = self.c
        chrome = CHROME[lang]
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.rect(x, y, width, height, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y, 4, height, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 6.8)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.drawString(x + 12, y + height - 14, f"MODULE {index:02d}")
        title_size, title_lines = fit_pdf_lines(topic["title"], FONT_BOLD, 9.4, 8.0, width - 28, 2)
        c.setFont(FONT_BOLD, title_size)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        title_y = y + height - 28
        for line in title_lines:
            c.drawString(x + 12, title_y, line)
            title_y -= round(title_size * MIN_LINE_RATIO, 2)

        inner_width = width - 24
        wrap_w = inner_width - 8
        tall = height >= 400
        if tall:
            meta_rows = [
                (chrome["role_short"], topic.get("role", "")),
                (chrome["flow_short"], topic.get("flow", "")),
                (chrome["boundary_short"], topic.get("engineeringBoundary", "")),
            ]
            meta_limit = 3
        else:
            meta_rows = []
            meta_limit = 2

        # 动态测量 meta 标签最大宽度，防止如英文 "Boundary" 与正文重叠
        meta_label_w = max((pdf_text_width(l, FONT_BOLD, 6.6) for l, _ in meta_rows), default=14.0)
        meta_text_x = x + 12 + meta_label_w + 4.5
        meta_wrap_w = inner_width - (meta_label_w + 4.5) - 4.0

        # 动态测量 outcome 结果栏标签宽度，防止与正文挨挤或重叠
        result_label_w = pdf_text_width(chrome["result"], FONT_BOLD, 6.8)
        result_text_x = x + 12 + result_label_w + 4.5
        result_wrap_w = (width - 16) - (12 + result_label_w + 4.5) - 4.0

        outcome_size, outcome_lines = fit_pdf_lines(topic["outcome"], FONT_REGULAR, 7.2, 6.4, result_wrap_w, 3)
        outcome_lead = line_leading(outcome_size) if tall else round(outcome_size * 1.4, 2)
        outcome_block = 8 + len(outcome_lines) * outcome_lead
        top = title_y - 4
        bottom = y + 6 + outcome_block
        available = top - bottom

        impl_size = 7.55 if tall else 6.9
        challenge_size = 7.35 if tall else 6.75
        meta_size = 6.9 if tall else 6.55
        impl_lines: list[list[str]] = []
        challenge_lines: list[list[str]] = []
        meta_wrapped: list[tuple[str, list[str]]] = []
        card_ratio = LINE_RATIO if tall else MIN_LINE_RATIO
        min_ratio = MIN_LINE_RATIO if tall else 1.46

        def pack(size_impl: float, size_challenge: float, size_meta: float) -> None:
            nonlocal impl_lines, challenge_lines, meta_wrapped
            impl_lines = [wrap_pdf_text(item, FONT_REGULAR, size_impl, wrap_w) for item in topic["implementation"]]
            challenge_lines = [wrap_pdf_text(item, FONT_REGULAR, size_challenge, wrap_w) for item in topic["challenges"]]
            meta_wrapped = [
                (label, wrap_pdf_text(text, FONT_REGULAR, size_meta, meta_wrap_w)[:meta_limit])
                for label, text in meta_rows
            ]

        def measure() -> float:
            meta_lead = round(meta_size * card_ratio, 2)
            impl_lead = round(impl_size * card_ratio, 2)
            chal_lead = round(challenge_size * card_ratio, 2)
            impl_gap = 5.6 if tall else 2.4
            chal_gap = 5.6 if tall else 2.4
            meta_gap = paragraph_gap(meta_size) * 0.55
            height_used = 0.0
            for _, lines in meta_wrapped:
                height_used += len(lines) * meta_lead + meta_gap
            height_used += 13
            for lines in impl_lines:
                height_used += len(lines) * impl_lead + impl_gap
            height_used += 13
            for lines in challenge_lines:
                height_used += len(lines) * chal_lead + chal_gap
            return height_used

        pack(impl_size, challenge_size, meta_size)
        while measure() > available - 6 and (impl_size > 6.0 or challenge_size > 6.0 or meta_size > 6.0):
            if impl_size > 6.0:
                impl_size = round(impl_size - 0.1, 2)
            if challenge_size > 6.0:
                challenge_size = round(challenge_size - 0.1, 2)
            if meta_size > 6.0:
                meta_size = round(meta_size - 0.08, 2)
            pack(impl_size, challenge_size, meta_size)
        while measure() > available - 4 and card_ratio > min_ratio:
            card_ratio = round(card_ratio - 0.02, 2)
        base_used = measure()
        if base_used > available + 1.6:
            impl_n = sum(len(x) for x in impl_lines)
            chal_n = sum(len(x) for x in challenge_lines)
            raise ValueError(
                f"PDF page {self.page_no} module card overflow: {topic['title']} "
                f"(used {base_used:.1f} avail {available:.1f} impl_lines={impl_n} chal_lines={chal_n} "
                f"size={impl_size}/{challenge_size} ratio={card_ratio} h={height})"
            )

        fill_reserve = 28 if tall else 0
        if base_used + fill_reserve > available:
            fill_reserve = max(0.0, available - base_used)
        leftover = max(0.0, available - fill_reserve - base_used)
        item_count = max(1, len(impl_lines) + len(challenge_lines))
        extra_gap = min(6.0 if tall else 3.2, leftover / item_count)
        meta_lead = round(meta_size * card_ratio, 2)
        impl_lead = round(impl_size * card_ratio, 2)
        chal_lead = round(challenge_size * card_ratio, 2)
        impl_gap = (5.6 if tall else 2.4) + extra_gap
        chal_gap = (5.6 if tall else 2.4) + extra_gap
        meta_gap = paragraph_gap(meta_size) * 0.55

        text_x = x + 24
        cursor = top
        for label, lines in meta_wrapped:
            c.setFont(FONT_BOLD, 6.6)
            c.setFillColor(HexColor(f"#{NAVY}"))
            c.drawString(x + 12, cursor, label)
            c.setFont(FONT_REGULAR, meta_size)
            c.setFillColor(HexColor(f"#{INK}"))
            line_y = cursor
            for line in lines:
                c.drawString(meta_text_x, line_y, line)
                line_y -= meta_lead
            cursor = line_y - meta_gap

        self.section_label(x + 12, cursor, chrome["implementation"])
        cursor -= 13
        for item_lines in impl_lines:
            c.setFillColor(HexColor(f"#{accent}"))
            c.circle(x + 16, cursor + 1.7, 1.7, fill=1, stroke=0)
            c.setFont(FONT_REGULAR, impl_size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in item_lines:
                c.drawString(text_x, cursor, line)
                cursor -= impl_lead
            cursor -= impl_gap

        c.setFont(FONT_BOLD, TYPE_SECTION)
        c.setFillColor(HexColor(f"#{NAVY}"))
        c.drawString(x + 12, cursor, chrome["challenges"])
        cursor -= 13
        for item_lines in challenge_lines:
            c.setFillColor(HexColor(f"#{ORANGE}"))
            c.circle(x + 16, cursor + 1.7, 1.7, fill=1, stroke=0)
            c.setFont(FONT_REGULAR, challenge_size)
            c.setFillColor(HexColor(f"#{INK}"))
            for line in item_lines:
                c.drawString(text_x, cursor, line)
                cursor -= chal_lead
            cursor -= chal_gap

        fill_bottom = y + 10 + outcome_block
        fill_h = max(0.0, cursor - 3 - fill_bottom)
        if fill_h >= 22:
            self.draw_card_fill(
                x + 12,
                fill_bottom,
                inner_width,
                fill_h,
                topic.get("background", ""),
                topic.get("stack", []),
                lang,
            )

        outcome_h = outcome_block + 4
        c.setFillColor(HexColor("#F3F7F8"))
        c.rect(x + 8, y + 6, width - 16, outcome_h, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 6.8)
        c.setFillColor(HexColor(f"#{TEAL}"))

        total_outcome_text_h = (len(outcome_lines) - 1) * outcome_lead + outcome_size * 0.72
        pad_outcome_y = (outcome_h - total_outcome_text_h) / 2
        text_y = y + 6 + outcome_h - pad_outcome_y - outcome_size * 0.72

        c.drawString(x + 12, text_y, chrome["result"])
        c.setFont(FONT_REGULAR, outcome_size)
        c.setFillColor(HexColor(f"#{INK}"))
        for line in outcome_lines:
            c.drawString(result_text_x, text_y, line)
            text_y -= outcome_lead

    def pipeline(self, y: float, flow: str, height: float = 32) -> None:
        nodes = [item.strip() for item in flow.split("→")]
        c = self.c
        c.setFillColor(HexColor(f"#{TEAL_SOFT}"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.5)
        c.roundRect(self.left, y, self.content_width, height, RADIUS, fill=1, stroke=1)
        gap = 7
        inset = 6
        node_h = height - 10
        node_y = y + 5
        node_width = (self.content_width - inset * 2 - gap * (len(nodes) - 1)) / len(nodes)
        for index, node in enumerate(nodes):
            x = self.left + inset + index * (node_width + gap)
            c.setFillColor(HexColor("#FFFFFF" if index % 2 == 0 else "#F7FBFC"))
            c.setStrokeColor(HexColor(f"#{TEAL}"))
            c.roundRect(x, node_y, node_width, node_h, 2, fill=1, stroke=1)
            node_size, lines = fit_pdf_lines(node, FONT_BOLD, 7.3, 5.8, node_width - 4, 1)
            c.setFont(FONT_BOLD, node_size)
            c.setFillColor(HexColor(f"#{NAVY}"))
            text_y = node_y + (node_h - 0.72 * node_size) / 2
            c.drawCentredString(x + node_width / 2, text_y, lines[0])
            if index < len(nodes) - 1:
                arrow_x = x + node_width + 1
                mid = node_y + node_h / 2
                c.setStrokeColor(HexColor(f"#{ORANGE}"))
                c.setLineWidth(1.1)
                c.line(arrow_x, mid, arrow_x + gap - 2, mid)
                c.line(arrow_x + gap - 5, mid + 2.4, arrow_x + gap - 2, mid)
                c.line(arrow_x + gap - 5, mid - 2.4, arrow_x + gap - 2, mid)

    def footer(self, lang: str) -> None:
        website = "http://cv.bookfree.online/"
        self.c.setFont(FONT_BOLD, 7.2)
        self.c.setFillColor(HexColor(f"#{TEAL}"))
        self.c.drawString(self.left, 17, f"3D PORTFOLIO: {website}")
        self.c.linkURL(website, (self.left, 10, self.left + 175, 25))
        self.c.setFont(FONT_REGULAR, 7.2)
        self.c.setFillColor(HexColor(f"#{MUTED}"))
        self.c.drawRightString(self.right, 17, f"CASEBOOK  /  {self.page_no:02d} OF {TOTAL_PAGES:02d}")


def draw_cover_chip_card(
    c: canvas.Canvas,
    x: float,
    y: float,
    box_w: float,
    box_h: float,
    text: str,
    accent: str,
) -> None:
    # 1. 绘制卡片背景矩形与边框
    c.setFillColor(HexColor("#102027"))
    c.setStrokeColor(HexColor("#264750"))
    c.setLineWidth(0.9)
    c.rect(x, y, box_w, box_h, fill=1, stroke=1)

    # 2. 左侧高亮竖条
    c.setFillColor(HexColor(f"#{accent}"))
    c.rect(x, y, 4.5, box_h, fill=1, stroke=0)

    # 3. 文本排版：优先单行，若超长自适应折行与字号拟合，垂直绝对居中，绝不溢出卡片
    avail_w = box_w - 20
    preferred_size = 8.4
    min_size = 6.2

    if pdf_text_width(text, FONT_BOLD, preferred_size) <= avail_w:
        size = preferred_size
        lines = [text]
    else:
        try:
            size, lines = fit_pdf_lines(text, FONT_BOLD, 8.0, 6.4, avail_w, 2)
        except ValueError:
            size, lines = fit_pdf_lines(text, FONT_BOLD, 7.0, min_size, avail_w, 3)

    lead = line_leading(size)
    text_block_h = (len(lines) - 1) * lead + size * 0.72
    pad_y = (box_h - text_block_h) / 2
    cur_y = y + box_h - pad_y - size * 0.72

    c.setFont(FONT_BOLD, size)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    for line in lines:
        c.drawString(x + 14, cur_y, line)
        cur_y -= lead


def draw_pdf_cover(c: canvas.Canvas, data: dict, lang: str) -> None:
    chrome = CHROME[lang]
    w, h = A4
    c.setFillColor(HexColor(f"#{GRAPHITE}"))
    c.rect(0, 0, w, h, fill=1, stroke=0)
    for x in range(40, 561, 48):
        c.setStrokeColor(HexColor("#10232B"))
        c.setLineWidth(0.35)
        c.line(x, 0, x, h)
    for y in range(30, 815, 48):
        c.setStrokeColor(HexColor("#10232B"))
        c.setLineWidth(0.35)
        c.line(0, y, w, y)

    # 左侧工业强调装饰条
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.rect(0, 0, 15, h, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(15, h - 145, 9, 95, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.rect(15, 200, 4, 180, fill=1, stroke=0)

    # 右上角工业管网/数据流拓扑 (Industrial Pipe & Data Topology)
    c.setStrokeColor(HexColor(f"#{CYAN}"))
    c.setLineWidth(1.8)
    path = c.beginPath()
    path.moveTo(345, 730)
    path.lineTo(455, 730)
    path.lineTo(455, 638)
    path.lineTo(541, 638)
    c.drawPath(path, stroke=1, fill=0)

    c.setStrokeColor(HexColor(f"#{ORANGE}"))
    c.setLineWidth(1.2)
    path2 = c.beginPath()
    path2.moveTo(435, 770)
    path2.lineTo(525, 770)
    path2.lineTo(525, 685)
    path2.lineTo(541, 685)
    c.drawPath(path2, stroke=1, fill=0)

    c.setStrokeColor(HexColor("#1A3B47"))
    c.setLineWidth(0.7)
    c.line(345, 670, 420, 670)
    c.line(420, 670, 455, 638)

    nodes = [
        (345, 730, 5, CYAN, "AIoT INGEST", 8, 4),
        (455, 730, 7, CYAN, "DMA TOPOLOGY", 8, 4),
        (455, 638, 6, TEAL, "CORE SERVICE", -58, -4),
        (541, 638, 8, CYAN, "AGENT GRAPH", -56, 5),
        (435, 770, 4.5, ORANGE, "KAFKA/FLINK", 7, 4),
        (525, 770, 5.5, ORANGE, "ELK INDEX", -42, 5),
        (541, 685, 6, ORANGE, "DISPATCH", -42, 5),
    ]
    for nx, ny, radius, col, lbl, tox, toy in nodes:
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        c.setStrokeColor(HexColor(f"#{col}"))
        c.setLineWidth(1.1)
        c.circle(nx, ny, radius, fill=1, stroke=1)
        c.setFillColor(HexColor(f"#{col}"))
        c.circle(nx, ny, 2.0, fill=1, stroke=0)
        c.setFont(FONT_REGULAR, 5.2)
        c.setFillColor(HexColor("#648894"))
        c.drawString(nx + tox, ny + toy, lbl)

    website = data["profile"].get("website", "http://cv.bookfree.online/")

    # 1. 顶部身份区域 (Top Identity Header)
    c.setFont(FONT_BOLD, 8.5)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawString(54, h - 55, "ENGINEERING CASEBOOK / 2026")
    c.setFont(FONT_REGULAR, 7.5)
    c.setFillColor(HexColor(f"#{TEAL}"))
    c.drawString(205, h - 55, f"[3D PORTFOLIO: {website}]")
    c.linkURL(website, (205, h - 60, 420, h - 48))

    c.setFont(FONT_BOLD, 36)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(54, h - 105, data["profile"]["name"])

    c.setFont(FONT_BOLD, 18)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(54, h - 134, data["profile"]["title"])

    sub_font_size = 9.5
    sub_text = f"{data['profile']['experience']}  ·  {chrome['cover_domains']}"
    while sub_font_size >= 7.8 and pdf_text_width(sub_text, FONT_REGULAR, sub_font_size) > 487:
        sub_font_size = round(sub_font_size - 0.2, 2)
    c.setFont(FONT_REGULAR, sub_font_size)
    c.setFillColor(HexColor("#9EB3BA"))
    c.drawString(54, h - 153, sub_text)

    # 分割线
    c.setStrokeColor(HexColor("#213E47"))
    c.setLineWidth(1.0)
    c.line(54, h - 170, 541, h - 170)
    c.setStrokeColor(HexColor(f"#{CYAN}"))
    c.setLineWidth(2.0)
    c.line(54, h - 170, 140, h - 170)

    # 2. 核心量化指标 (Core Scale Metric)
    c.setFont(FONT_BOLD, 46)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawString(54, h - 235, "10w+")
    c.setFont(FONT_BOLD, 10.5)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(56, h - 256, chrome["cover_water"])

    # 3. 核心亮点工业线框矩阵 (Highlight Chip Boxes Matrix - 2列多行)
    box_w = 236
    box_h = 42
    col_gap = 15
    row_gap = 9
    start_y = h - 316

    accent_colors = [TEAL, CYAN, TEAL, ORANGE, CYAN, TEAL, ORANGE, CYAN, TEAL, ORANGE]

    for index, item in enumerate(data["highlights"]):
        col = index % 2
        row = index // 2
        x = 54 + col * (box_w + col_gap)
        y = start_y - row * (box_h + row_gap)
        accent = accent_colors[index % len(accent_colors)]
        draw_cover_chip_card(c, x, y, box_w, box_h, item, accent)

    # 4. 中间件基础设施矩阵 (Middleware Fabric - 2行2列)
    middleware_items = (
        ("Redis Cluster", TEAL),
        ("RabbitMQ", ORANGE),
        ("Nacos", CYAN),
        ("Sentinel", TEAL),
    )
    middleware_start_y = 275

    for index, (technology, accent) in enumerate(middleware_items):
        col = index % 2
        row = index // 2
        x = 54 + col * (box_w + col_gap)
        y = middleware_start_y - row * (box_h + row_gap)
        draw_cover_chip_card(c, x, y, box_w, box_h, technology, accent)

    # 5. 职业定位与工程准则终端框 (Engineering Statement Terminal)
    stmt_y = 120
    stmt_h = 100
    c.setFillColor(HexColor("#0A1820"))
    c.setStrokeColor(HexColor("#1F424E"))
    c.roundRect(54, stmt_y, 487, stmt_h, 3, fill=1, stroke=1)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(54, stmt_y + stmt_h - 3, 487, 3, fill=1, stroke=0)

    c.setFont(FONT_BOLD, 7.8)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.drawString(68, stmt_y + stmt_h - 18, chrome["cover_statement"])

    c.setFont(FONT_REGULAR, 8.8)
    c.setFillColor(HexColor("#CCDCE0"))
    summary_lines = wrap_pdf_text(data["profile"]["summary"], FONT_REGULAR, 8.8, 458)
    for s_idx, s_line in enumerate(summary_lines):
        c.drawString(68, stmt_y + stmt_h - 35 - s_idx * 14, s_line)

    c.setFont(FONT_REGULAR, 7.6)
    c.setFillColor(HexColor("#7E9BA4"))
    c.drawString(68, stmt_y + 16, chrome["cover_statement_body"])

    # 6. 底部联络与状态栏 (Footer & Contact Terminal)
    c.setStrokeColor(HexColor(f"#{TEAL}"))
    c.setLineWidth(0.8)
    c.line(54, 88, 541, 88)

    c.setFont(FONT_BOLD, 8.2)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(54, 69, f"TEL: {data['profile']['phone']}")
    c.drawString(175, 69, f"EMAIL: {data['profile']['email']}")
    c.drawString(380, 69, chrome["cover_base"])

    c.setFont(FONT_BOLD, 8.2)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(54, 51, f"3D PORTFOLIO SITE: {website}")
    c.linkURL(website, (54, 46, 320, 58))

    c.setFont(FONT_REGULAR, 7.0)
    c.setFillColor(HexColor("#5D7881"))
    c.drawString(54, 28, chrome["cover_confidential"].format(name=data["profile"]["name"]))
    c.drawRightString(541, 28, f"01 / {TOTAL_PAGES:02d} · JAVA BACKEND · AIOT · GIS · AGENT")

    c.setFont(FONT_REGULAR, 8.8)
    c.setFillColor(HexColor("#CCDCE0"))
    summary_lines = wrap_pdf_text(data["profile"]["summary"], FONT_REGULAR, 8.8, 458)
    for s_idx, s_line in enumerate(summary_lines):
        c.drawString(68, stmt_y + stmt_h - 35 - s_idx * 14, s_line)

    c.setFont(FONT_REGULAR, 7.6)
    c.setFillColor(HexColor("#7E9BA4"))
    c.drawString(68, stmt_y + 16, chrome["cover_statement_body"])

    # 6. 底部联络与状态栏 (Footer & Contact Terminal)
    c.setStrokeColor(HexColor(f"#{TEAL}"))
    c.setLineWidth(0.8)
    c.line(54, 88, 541, 88)

    c.setFont(FONT_BOLD, 8.2)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    c.drawString(54, 69, f"TEL: {data['profile']['phone']}")
    c.drawString(175, 69, f"EMAIL: {data['profile']['email']}")
    c.drawString(380, 69, chrome["cover_base"])

    c.setFont(FONT_BOLD, 8.2)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(54, 51, f"3D PORTFOLIO SITE: {website}")
    c.linkURL(website, (54, 46, 320, 58))

    c.setFont(FONT_REGULAR, 7.0)
    c.setFillColor(HexColor("#5D7881"))
    c.drawString(54, 28, chrome["cover_confidential"].format(name=data["profile"]["name"]))
    c.drawRightString(541, 28, f"01 / {TOTAL_PAGES:02d} · JAVA BACKEND · AIOT · GIS · AGENT")


def draw_pdf_overview(c: canvas.Canvas, data: dict, lang: str) -> None:
    chrome = CHROME[lang]
    website = data["profile"].get("website", "http://cv.bookfree.online/")
    page = PdfPage(c, 2, "PROFILE / CAPABILITY MATRIX", data["profile"]["name"])
    page.title(chrome["overview"], chrome["overview_kicker"].format(website=website))
    page.info_box(page.left, 694, page.content_width, 48, chrome["overview_position"], data["profile"]["summary"], ORANGE)
    page.section_label(page.left, 678, chrome["overview_strengths"], "A1")
    box_gap = 8
    box_width = (page.content_width - box_gap) / 2
    strength_h = 54
    for index, strength in enumerate(data["strengths"]):
        x = page.left + (index % 2) * (box_width + box_gap)
        y = 616 - (index // 2) * (strength_h + 7)
        page.info_box(x, y, box_width, strength_h, strength["title"], strength["evidence"], TEAL if index < 3 else ORANGE)

    stack: list[str] = []
    for project in data["projects"]:
        for item in project.get("stack", []):
            if item not in stack:
                stack.append(item)
        for topic in project["topics"]:
            for item in topic["stack"]:
                if item not in stack:
                    stack.append(item)

    page.section_label(page.left, 538, chrome["overview_experience"], "A2")
    stack_h = 48
    stack_y = 62
    area_top = 526
    area_bottom = stack_y + stack_h + 8
    experiences = data["experiences"]
    gap = 8
    card_h = (area_top - area_bottom - gap * (len(experiences) - 1)) / len(experiences)
    for index, exp in enumerate(experiences):
        y = area_top - (index + 1) * card_h - index * gap
        page.experience_card(
            page.left,
            y,
            page.content_width,
            card_h,
            exp,
            ORANGE if index == 0 else TEAL,
        )
    page.info_box(page.left, stack_y, page.content_width, stack_h, chrome["overview_stack"], " · ".join(stack[:24]), ORANGE)
    page.footer(lang)


def draw_outcome_bar(page: PdfPage, text: str, lang: str) -> None:
    c = page.c
    chrome = CHROME[lang]
    bar_y = 56
    bar_h = 50
    c.setFillColor(HexColor(f"#{GRAPHITE}"))
    c.rect(page.left, bar_y, page.content_width, bar_h, fill=1, stroke=0)
    c.setFillColor(HexColor(f"#{ORANGE}"))
    c.rect(page.left, bar_y, 5, bar_h, fill=1, stroke=0)

    label_size = 7.8
    text_size = 8.45
    lead = line_leading(text_size, wide=True)
    lines = wrap_pdf_text(text, FONT_REGULAR, text_size, page.content_width - 28)[:2]

    label_gap = 4.5
    total_h = label_size * 0.72 + label_gap + (len(lines) - 1) * lead + text_size * 0.72
    pad_top = (bar_h - total_h) / 2

    label_y = bar_y + bar_h - pad_top - label_size * 0.72
    c.setFont(FONT_BOLD, label_size)
    c.setFillColor(HexColor(f"#{CYAN}"))
    c.drawString(page.left + 14, label_y, chrome["outcome"])

    text_y = label_y - label_gap - text_size * 0.72
    c.setFont(FONT_REGULAR, text_size)
    c.setFillColor(HexColor(f"#{COLD_WHITE}"))
    for line in lines:
        c.drawString(page.left + 14, text_y, line)
        text_y -= lead


def draw_stack_line(page: PdfPage, stack: list[str]) -> None:
    page.c.setFont(FONT_BOLD, 7.4)
    page.c.setFillColor(HexColor(f"#{TEAL}"))
    page.c.drawString(page.left, 41, "TECH MATRIX")
    page.wrapped("  /  ".join(stack), page.left + 72, 41, page.content_width - 72, size=7.15, leading=line_leading(7.15, wide=True), color=MUTED, max_lines=2)


def draw_pdf_system_overview(c: canvas.Canvas, page_no: int, project: dict, lang: str, name: str) -> None:
    chrome = CHROME[lang]
    page = PdfPage(c, page_no, f"{project['name'].upper()}  /  {chrome['system_arch']}", name)
    page.title(project["name"], f"{project['company']}  |  {project['period']}")
    page.info_box(page.left, 694, page.content_width, 48, chrome["system_position"], project["summary"], ORANGE)

    page.section_label(page.left, 678, chrome["system_background"], "01")
    page.labeled_rows_box(
        page.left,
        566,
        page.content_width,
        104,
        [
            (chrome["business"], project["businessContext"], TEAL),
            (chrome["pain"], join_items(project["painPoints"], lang), ORANGE),
            (chrome["goals"], join_items(project["buildGoals"], lang), TEAL),
        ],
    )

    page.section_label(page.left, 550, chrome["heading_flow"], "02")
    page.pipeline(510, project["flow"], 32)

    page.section_label(page.left, 494, chrome["heading_duty"], "03")
    page.labeled_rows_box(
        page.left,
        432,
        page.content_width,
        54,
        [
            (chrome["role"], project["role"], TEAL),
            (chrome["boundary"], project["engineeringBoundary"], ORANGE),
        ],
        preferred_size=8.1,
        min_size=6.4,
    )

    page.section_label(page.left, 416, chrome["modules"], "04")
    gap = 8
    card_width = (page.content_width - gap * (len(project["topics"]) - 1)) / len(project["topics"])
    accents = [TEAL, CYAN, ORANGE]
    card_y = 116
    card_h = 292
    for index, topic in enumerate(project["topics"]):
        x = page.left + index * (card_width + gap)
        y = card_y
        height = card_h
        c.setFillColor(HexColor("#FFFFFF"))
        c.setStrokeColor(HexColor(f"#{RULE}"))
        c.setLineWidth(0.6)
        c.rect(x, y, card_width, height, fill=1, stroke=1)
        accent = accents[index % len(accents)]
        c.setFillColor(HexColor(f"#{accent}"))
        c.rect(x, y + height - 4, card_width, 4, fill=1, stroke=0)
        c.setFont(FONT_BOLD, 7)
        c.setFillColor(HexColor(f"#{ORANGE}"))
        c.drawString(x + 10, y + height - 16, f"MODULE {index + 1:02d}")
        c.setFont(FONT_BOLD, 9.6)
        c.setFillColor(HexColor(f"#{GRAPHITE}"))
        title_size, title_lines = fit_pdf_lines(topic["title"], FONT_BOLD, 9.6, 8.0, card_width - 20, 2)
        title_y = y + height - 32
        c.setFont(FONT_BOLD, title_size)
        for line in title_lines:
            c.drawString(x + 10, title_y, line)
            title_y -= round(title_size * MIN_LINE_RATIO, 2)
        after_bg = page.wrapped(topic["background"], x + 10, title_y - 8, card_width - 20, size=7.8, leading=line_leading(7.8), max_lines=6, min_size=6.8)
        after_role = page.wrapped(f"{chrome['role_short']}  {topic['role']}", x + 10, after_bg - 8, card_width - 20, size=7.3, leading=line_leading(7.3), color=MUTED, max_lines=4, min_size=6.6)
        if topic.get("stack"):
            page.draw_stack_chips(x + 10, y + 10, card_width - 20, max(24.0, after_role - y - 16), topic["stack"])

    draw_outcome_bar(page, project["outcome"], lang)
    draw_stack_line(page, project["stack"])
    page.footer(lang)


def draw_pdf_system_modules(c: canvas.Canvas, page_no: int, project: dict, lang: str, name: str) -> None:
    chrome = CHROME[lang]
    page = PdfPage(c, page_no, f"{project['name'].upper()}  /  {chrome['practice']}", name)
    page.title(f"{project['name']} · {chrome['practice']}", f"{project['company']}  |  {project['period']}")
    gap = 8
    card_width = (page.content_width - gap * (len(project["topics"]) - 1)) / len(project["topics"])
    accents = [TEAL, CYAN, ORANGE]
    card_y = 116
    card_h = 626
    for index, topic in enumerate(project["topics"]):
        page.module_card(
            page.left + index * (card_width + gap),
            card_y,
            card_width,
            card_h,
            index + 1,
            topic,
            accents[index % len(accents)],
            lang,
        )
    draw_outcome_bar(page, project["outcome"], lang)
    draw_stack_line(page, project["stack"])
    page.footer(lang)


def draw_pdf_system_onepager(c: canvas.Canvas, page_no: int, project: dict, lang: str, name: str) -> None:
    chrome = CHROME[lang]
    page = PdfPage(c, page_no, f"{project['name'].upper()}  /  {chrome['system_practice']}", name)
    page.title(project["name"], f"{project['company']}  |  {project['period']}")
    page.info_box(page.left, 694, page.content_width, 48, chrome["system_position"], project["summary"], ORANGE)

    page.section_label(page.left, 678, chrome["system_background"], "01")
    page.labeled_rows_box(
        page.left,
        566,
        page.content_width,
        104,
        [
            (chrome["business"], project["businessContext"], TEAL),
            (chrome["pain"], join_items(project["painPoints"], lang), ORANGE),
            (chrome["goals"], join_items(project["buildGoals"], lang), TEAL),
        ],
        preferred_size=7.6,
        min_size=6.2,
    )

    page.section_label(page.left, 550, chrome["heading_flow"], "02")
    page.pipeline(510, project["flow"], 32)

    page.section_label(page.left, 494, chrome["heading_duty"], "03")
    page.labeled_rows_box(
        page.left,
        432,
        page.content_width,
        54,
        [
            (chrome["role"], project["role"], TEAL),
            (chrome["boundary"], project["engineeringBoundary"], ORANGE),
        ],
        preferred_size=7.6,
        min_size=6.2,
    )

    page.section_label(page.left, 416, chrome["impl_challenges"], "04")
    gap = 8
    card_y, card_h = 116, 292
    topics = project["topics"]
    if len(topics) == 1:
        topic = topics[0]
        implementation_width, challenge_width = choose_bullet_columns(
            page.content_width,
            gap,
            topic["implementation"],
            topic["challenges"],
        )
        page.bullet_box(page.left, card_y, implementation_width, card_h, f"{chrome['module']} 01  {topic['title']} · {chrome['implementation']}", topic["implementation"], TEAL)
        page.bullet_box(page.left + implementation_width + gap, card_y, challenge_width, card_h, chrome["challenges"], topic["challenges"], ORANGE)
    else:
        card_width = (page.content_width - gap * (len(topics) - 1)) / len(topics)
        accents = [TEAL, CYAN, ORANGE]
        for index, topic in enumerate(topics):
            page.module_card(
                page.left + index * (card_width + gap),
                card_y,
                card_width,
                card_h,
                index + 1,
                topic,
                accents[index % len(accents)],
                lang,
            )

    draw_outcome_bar(page, project["outcome"], lang)
    draw_stack_line(page, project["stack"])
    page.footer(lang)


def build_pdf(data: dict, lang: str, out_path: Path) -> None:
    chrome = CHROME[lang]
    name = data["profile"]["name"]
    register_pdf_fonts()
    c = canvas.Canvas(str(out_path), pagesize=A4, pageCompression=1)
    c.setTitle(f"{name} - {data['profile']['title']}")
    c.setAuthor(name)
    c.setSubject(chrome["pdf_subject"])
    draw_pdf_cover(c, data, lang)
    c.showPage()
    draw_pdf_overview(c, data, lang)

    projects = data["projects"]
    c.showPage()
    draw_pdf_system_overview(c, 3, projects[0], lang, name)
    c.showPage()
    draw_pdf_system_modules(c, 4, projects[0], lang, name)
    for page_no, project in enumerate(projects[1:], start=5):
        c.showPage()
        draw_pdf_system_onepager(c, page_no, project, lang, name)
    c.save()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--lang", choices=("zh", "en", "all"), default="all")
    args = parser.parse_args()
    langs = ("zh", "en") if args.lang == "all" else (args.lang,)
    public_resume_dir = ROOT / "web" / "public" / "resume"
    public_resume_dir.mkdir(parents=True, exist_ok=True)
    for lang in langs:
        data = load_data(lang)
        bundle = resume_bundle(lang)
        build_docx(data, lang, bundle["docx"])
        build_pdf(data, lang, bundle["pdf"])
        shutil.copy2(bundle["pdf"], public_resume_dir / bundle["pdf"].name)
        shutil.copy2(bundle["docx"], public_resume_dir / bundle["docx"].name)
        print(bundle["docx"])
        print(bundle["pdf"])
    print(f"Copied to {public_resume_dir}")


if __name__ == "__main__":
    main()
